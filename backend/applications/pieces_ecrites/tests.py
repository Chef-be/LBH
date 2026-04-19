from io import BytesIO
from unittest.mock import patch

from django.test import TestCase
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document as DocumentWord
from openpyxl import Workbook, load_workbook
from rest_framework import status
from rest_framework.test import APIClient

from applications.comptes.models import Utilisateur
from applications.bibliotheque.models import LignePrixBibliotheque
from applications.documents.models import Document, TypeDocument
from applications.organisations.models import Organisation
from applications.economie.models import EtudeEconomique, LignePrix
from applications.pieces_ecrites.models import ArticleCCTP, ModeleDocument, PieceEcrite
from applications.pieces_ecrites.office import creer_jeton_wopi_modele
from applications.pieces_ecrites.services import (
    construire_donnees_fusion_piece,
    exporter_piece_ecrite,
    generer_contenu_piece_depuis_articles,
    generer_piece_depuis_modele,
    importer_fichier_word_en_html,
)
from applications.projets.models import Projet


class FauxFluxHttp:
    def __init__(self, contenu: bytes):
        self.contenu = contenu

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, traceback):
        return False

    def read(self):
        return self.contenu


class ServicesPiecesEcritesTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.organisation = Organisation.objects.create(
            code="LBH",
            nom="LBH",
            type_organisation="bureau_etudes",
        )
        self.maitre_ouvrage = Organisation.objects.create(
            code="MOA",
            nom="Commune de Mamoudzou",
            type_organisation="maitre_ouvrage",
        )
        self.responsable = Utilisateur.objects.create_user(
            courriel="redacteur@example.com",
            password="motdepasse123",
            prenom="Anne",
            nom="Rédactrice",
            organisation=self.organisation,
        )
        self.client.force_authenticate(self.responsable)
        self.projet = Projet.objects.create(
            reference="2026-0002",
            intitule="Construction d'un pôle administratif",
            organisation=self.organisation,
            maitre_ouvrage=self.maitre_ouvrage,
            responsable=self.responsable,
        )
        self.modele = ModeleDocument.objects.create(
            code="CCTP_STANDARD_TEST",
            libelle="CCTP standard",
            type_document="cctp",
            variables_fusion=[
                {"nom": "nom_projet", "description": "Nom du projet"},
                {"nom": "reference_projet", "description": "Référence du projet"},
                {"nom": "maitre_ouvrage", "description": "Maître d'ouvrage"},
            ],
        )
        self.piece = PieceEcrite.objects.create(
            projet=self.projet,
            modele=self.modele,
            intitule="CCTP Lot Gros Œuvre",
            redacteur=self.responsable,
        )
        ArticleCCTP.objects.create(
            piece_ecrite=self.piece,
            chapitre="1",
            numero_article="1.1",
            intitule="Béton de propreté",
            corps_article="Le béton de propreté sera dosé conformément aux prescriptions du marché.",
            normes_applicables="NF EN 206",
        )
        self.etude = EtudeEconomique.objects.create(
            projet=self.projet,
            intitule="Étude économique de référence",
            cree_par=self.responsable,
        )
        LignePrix.objects.create(
            etude=self.etude,
            numero_ordre=1,
            code="VRD-01",
            designation="Terrassements généraux",
            unite="m3",
            quantite_prevue="12.500",
            temps_main_oeuvre="1.2500",
            cout_horaire_mo="42.0000",
            prix_vente_unitaire="95.0000",
        )
        self.ligne_bibliotheque = LignePrixBibliotheque.objects.create(
            code="LBH-PRI-TEST-001",
            famille="Gros œuvre",
            lot="Lot 03",
            designation_longue="Béton de propreté compris fourniture et mise en œuvre.",
            designation_courte="Béton de propreté",
            unite="m3",
            prescriptions_techniques="Le béton sera coulé sur support propre et préparé.",
            criteres_metre="Le métré est réalisé au mètre cube réellement mis en œuvre.",
            normes_applicables=["NF EN 206", "DTU 21"],
            phases_execution=["Préparation du support", "Coulage", "Protection et cure"],
            dechets_generes=["Déchets inertes de béton"],
            cahier_des_charges_structure=[
                {"titre": "Préparation", "contenu": "Support débarrassé des éléments non adhérents."},
                {"titre": "Exécution", "contenu": "Coulage et dressage selon les règles de l'art."},
            ],
            source="Référentiel test",
        )

    def test_generation_piece_depuis_articles(self):
        contenu = generer_contenu_piece_depuis_articles(self.piece)

        self.assertIn("CCTP Lot Gros Œuvre", contenu)
        self.assertIn("2026-0002", contenu)
        self.assertIn("Béton de propreté", contenu)
        self.assertIn("NF EN 206", contenu)

    def test_export_piece_ecrite_docx_et_pdf(self):
        contenu_docx, type_docx, nom_docx = exporter_piece_ecrite(self.piece, "docx")
        contenu_pdf, type_pdf, nom_pdf = exporter_piece_ecrite(self.piece, "pdf")

        self.assertTrue(contenu_docx.startswith(b"PK"))
        self.assertTrue(contenu_pdf.startswith(b"%PDF"))
        self.assertEqual(type_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertEqual(type_pdf, "application/pdf")
        self.assertTrue(nom_docx.endswith(".docx"))
        self.assertTrue(nom_pdf.endswith(".pdf"))
        self.piece.refresh_from_db()
        self.assertIsNotNone(self.piece.document_ged)
        self.assertTrue(Document.objects.filter(pk=self.piece.document_ged_id).exists())
        self.assertEqual(self.piece.document_ged.projet_id, self.projet.id)

    def test_generation_depuis_modele_et_variables(self):
        self.piece.variables_personnalisees = {"maitre_ouvrage": "Commune de Mamoudzou"}
        self.piece.save(update_fields=["variables_personnalisees"])

        generer_piece_depuis_modele(self.piece)
        self.piece.refresh_from_db()
        fusion = construire_donnees_fusion_piece(self.piece)

        self.assertEqual(fusion["maitre_ouvrage"], "Commune de Mamoudzou")
        self.assertIn("CCTP standard", self.piece.contenu_html)
        self.assertIn("Commune de Mamoudzou", self.piece.contenu_html)
        self.assertIn(self.projet.reference, self.piece.contenu_html)

    def test_generation_depuis_contenu_modele_html(self):
        self.modele.contenu_modele_html = (
            "<h1>{piece_intitule}</h1>"
            "<p><strong>Projet :</strong> {reference_projet}</p>"
            "<p><strong>Maître d'ouvrage :</strong> {maitre_ouvrage}</p>"
        )
        self.modele.save(update_fields=["contenu_modele_html"])
        self.piece.variables_personnalisees = {"maitre_ouvrage": "Ville de Mamoudzou"}
        self.piece.save(update_fields=["variables_personnalisees"])

        generer_piece_depuis_modele(self.piece)
        self.piece.refresh_from_db()

        self.assertIn("CCTP Lot Gros Œuvre", self.piece.contenu_html)
        self.assertIn(self.projet.reference, self.piece.contenu_html)
        self.assertIn("Ville de Mamoudzou", self.piece.contenu_html)
        self.assertNotIn("Données de fusion", self.piece.contenu_html)

    def test_export_docx_depuis_gabarit_modele(self):
        document = DocumentWord()
        document.add_paragraph("Projet : {nom_projet}")
        document.add_paragraph("Référence : {reference_projet}")
        document.add_paragraph("Contenu : {contenu_principal}")
        flux = BytesIO()
        document.save(flux)
        flux.seek(0)

        self.modele.gabarit = SimpleUploadedFile(
            "gabarit.docx",
            flux.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.modele.save(update_fields=["gabarit"])

        self.piece.variables_personnalisees = {"nom_projet": "Pôle administratif central"}
        self.piece.contenu_html = "<p>Contenu de démonstration</p>"
        self.piece.save(update_fields=["variables_personnalisees", "contenu_html"])

        contenu_docx, _, _ = exporter_piece_ecrite(self.piece, "docx")
        document_exporte = DocumentWord(BytesIO(contenu_docx))
        texte = "\n".join(paragraphe.text for paragraphe in document_exporte.paragraphs)

        self.assertIn("Pôle administratif central", texte)
        self.assertIn(self.projet.reference, texte)
        self.assertIn("Contenu de démonstration", texte)

    def test_export_xlsx_depuis_gabarit_modele_tableur(self):
        modele_bpu = ModeleDocument.objects.create(
            code="BPU_STANDARD_TEST",
            libelle="BPU standard",
            type_document="bpu",
            variables_fusion=[
                {"nom": "nom_projet", "description": "Nom du projet"},
                {"nom": "reference_projet", "description": "Référence du projet"},
            ],
        )
        piece_bpu = PieceEcrite.objects.create(
            projet=self.projet,
            modele=modele_bpu,
            intitule="BPU Lot VRD",
            redacteur=self.responsable,
        )

        classeur = Workbook()
        feuille = classeur.active
        feuille["A1"] = "{reference_projet}"
        feuille["B1"] = "{nom_projet}"
        flux = BytesIO()
        classeur.save(flux)
        flux.seek(0)

        modele_bpu.gabarit = SimpleUploadedFile(
            "gabarit.xlsx",
            flux.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        modele_bpu.save(update_fields=["gabarit"])

        piece_bpu.variables_personnalisees = {"nom_projet": "Pôle administratif central"}
        piece_bpu.save(update_fields=["variables_personnalisees"])

        contenu_xlsx, type_xlsx, nom_xlsx = exporter_piece_ecrite(piece_bpu, "xlsx")
        classeur_exporte = load_workbook(BytesIO(contenu_xlsx))
        feuille_exportee = classeur_exporte.active

        self.assertEqual(type_xlsx, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.assertTrue(nom_xlsx.endswith(".xlsx"))
        self.assertEqual(feuille_exportee["A1"].value, self.projet.reference)
        self.assertEqual(feuille_exportee["B1"].value, "Pôle administratif central")

    def test_import_word_en_html_conserve_structure(self):
        document = DocumentWord()
        document.add_heading("Titre importé", level=1)
        paragraphe = document.add_paragraph()
        paragraphe.add_run("Texte ")
        run = paragraphe.add_run("important")
        run.bold = True
        flux = BytesIO()
        document.save(flux)
        flux.seek(0)

        resultat = importer_fichier_word_en_html(
            SimpleUploadedFile(
                "import.docx",
                flux.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        )

        self.assertIn("Titre importé", str(resultat["html"]))
        self.assertIn("<strong>important</strong>", str(resultat["html"]))

    def test_generation_modele_planning_depuis_lignes_economiques(self):
        modele_planning = ModeleDocument.objects.create(
            code="PLANNING_TEST",
            libelle="Planning depuis étude",
            type_document="planning_taches",
        )
        piece_planning = PieceEcrite.objects.create(
            projet=self.projet,
            modele=modele_planning,
            intitule="Planning prévisionnel lot VRD",
            redacteur=self.responsable,
        )

        generer_piece_depuis_modele(piece_planning)
        piece_planning.refresh_from_db()

        self.assertIn("Planning prévisionnel des tâches", piece_planning.contenu_html)
        self.assertIn("Terrassements généraux", piece_planning.contenu_html)
        self.assertIn("jour(s)", piece_planning.contenu_html)

    def test_generation_memoire_technique_prend_en_compte_client_public(self):
        modele_memoire = ModeleDocument.objects.create(
            code="MEMOIRE_TEST",
            libelle="Mémoire technique",
            type_document="memoire_technique",
        )
        piece_memoire = PieceEcrite.objects.create(
            projet=self.projet,
            modele=modele_memoire,
            intitule="Mémoire technique de candidature",
            redacteur=self.responsable,
        )

        generer_piece_depuis_modele(piece_memoire)
        piece_memoire.refresh_from_db()

        self.assertIn("commande publique", piece_memoire.contenu_html)
        self.assertIn("Compréhension du besoin", piece_memoire.contenu_html)
        self.assertIn("Pondération indicative", piece_memoire.contenu_html)

    def test_proposition_article_cctp_genere_un_brouillon_depuis_la_bibliotheque(self):
        resultat = self.client.post(
            f"/api/pieces-ecrites/{self.piece.id}/proposition-cctp/",
            {
                "chapitre": "2",
                "numero_article": "2.1",
                "intitule": "",
                "ligne_prix_reference": str(self.ligne_bibliotheque.id),
                "niveau_detail": "detaille",
                "inclure_mise_en_oeuvre": True,
                "inclure_controles": True,
                "inclure_dechets": True,
            },
            format="json",
        )

        self.assertEqual(resultat.status_code, status.HTTP_200_OK, resultat.data)
        article = resultat.data["article"]
        self.assertEqual(article["ligne_prix_reference"], str(self.ligne_bibliotheque.id))
        self.assertEqual(article["intitule"], "Béton de propreté")
        self.assertIn("Gestion des déchets", article["corps_article"])
        self.assertIn("NF EN 206", article["normes_applicables"])


class ApiModelesDocumentsTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.organisation = Organisation.objects.create(
            code="ORG-PE",
            nom="Organisation pièces écrites",
            type_organisation="bureau_etudes",
        )
        self.admin = Utilisateur.objects.create_user(
            courriel="admin-pieces@example.com",
            password="motdepasse123",
            prenom="Admin",
            nom="Pièces écrites",
            organisation=self.organisation,
            est_staff=True,
            est_super_admin=True,
        )
        self.client.force_authenticate(self.admin)
        self.modele = ModeleDocument.objects.create(
            code="MODELE_UPLOAD_TEST",
            libelle="Modèle avec gabarit",
            type_document="cctp",
        )

    def test_mise_a_jour_modele_accepte_un_patch_multipart_avec_fichier(self):
        reponse = self.client.patch(
            f"/api/pieces-ecrites/modeles/{self.modele.id}/",
            {
                "libelle": "Modèle avec gabarit mis à jour",
                "gabarit": SimpleUploadedFile(
                    "gabarit.docx",
                    b"contenu-docx-simule",
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
            format="multipart",
        )

        self.assertEqual(reponse.status_code, status.HTTP_200_OK, reponse.data)
        self.modele.refresh_from_db()
        self.assertEqual(self.modele.libelle, "Modèle avec gabarit mis à jour")
        self.assertTrue(bool(self.modele.gabarit))
        self.assertIn("gabarits/", self.modele.gabarit.name)

    @patch("applications.pieces_ecrites.office.urllib_request.urlopen")
    def test_session_bureautique_retourne_configuration_collabora(self, mock_urlopen):
        mock_urlopen.return_value = FauxFluxHttp(
            b"""
            <wopi-discovery>
              <net-zone name="external-https">
                <app name="writer">
                  <action ext="docx" name="edit" urlsrc="https://office.example.test/browser/cool.html?" />
                </app>
              </net-zone>
            </wopi-discovery>
            """
        )

        reponse = self.client.post(f"/api/pieces-ecrites/modeles/{self.modele.id}/session-bureautique/", {})

        self.assertEqual(reponse.status_code, status.HTTP_200_OK, reponse.data)
        self.assertEqual(reponse.data["type_bureautique"], "texte")
        self.assertEqual(reponse.data["extension"], ".docx")
        self.assertTrue(reponse.data["access_token"])
        self.assertEqual(reponse.data["access_token_ttl"], 28800000)
        self.assertIn("WOPISrc=", reponse.data["url_editeur"])
        self.assertNotIn("access_token=", reponse.data["url_editeur"])

    def test_wopi_permet_de_lire_et_sauvegarder_un_gabarit(self):
        jeton = creer_jeton_wopi_modele(self.modele, self.admin)
        url_fichier = f"/api/pieces-ecrites/wopi/modeles/{self.modele.id}/?access_token={jeton}"
        url_contenu = f"/api/pieces-ecrites/wopi/modeles/{self.modele.id}/contents?access_token={jeton}"

        reponse_info = self.client.get(url_fichier)
        self.assertEqual(reponse_info.status_code, status.HTTP_200_OK, reponse_info.data)
        self.assertEqual(reponse_info.data["SupportsUpdate"], True)
        self.assertEqual(reponse_info.data["SupportsGetLock"], True)
        self.assertTrue(reponse_info.data["BaseFileName"].endswith(".docx"))

        reponse_verrou = self.client.post(
            url_fichier,
            data=b"",
            content_type="application/octet-stream",
            HTTP_X_WOPI_OVERRIDE="LOCK",
            HTTP_X_WOPI_LOCK="verrou-modele-1",
        )
        self.assertEqual(reponse_verrou.status_code, status.HTTP_200_OK)

        document = DocumentWord()
        document.add_paragraph("Projet : {nom_projet}")
        flux = BytesIO()
        document.save(flux)
        flux.seek(0)

        reponse_sauvegarde = self.client.post(
            url_contenu,
            data=flux.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            HTTP_X_WOPI_OVERRIDE="PUT",
            HTTP_X_WOPI_LOCK="verrou-modele-1",
        )
        self.assertEqual(reponse_sauvegarde.status_code, status.HTTP_200_OK)

        reponse_contenu = self.client.get(url_contenu)
        self.assertEqual(reponse_contenu.status_code, status.HTTP_200_OK)
        document_enregistre = DocumentWord(BytesIO(reponse_contenu.content))
        self.assertIn("Projet : {nom_projet}", "\n".join(p.text for p in document_enregistre.paragraphs))

    @patch("applications.documents.office.urllib_request.urlopen")
    def test_session_bureautique_piece_reutilise_le_document_ged_associe(self, mock_urlopen):
        mock_urlopen.return_value = FauxFluxHttp(
            b"""
            <wopi-discovery>
              <net-zone name="external-https">
                <app name="writer">
                  <action ext="docx" name="edit" urlsrc="https://office.example.test/browser/cool.html?" />
                </app>
              </net-zone>
            </wopi-discovery>
            """
        )
        projet = Projet.objects.create(
            reference="2026-0501",
            intitule="Projet GED",
            organisation=self.organisation,
            responsable=self.admin,
        )
        type_document = TypeDocument.objects.create(code="CCTP-T", libelle="CCTP test")
        document = Document.objects.create(
            reference="DOC-001",
            intitule="Document GED lié",
            type_document=type_document,
            projet=projet,
            nom_fichier_origine="document-ged-lie.docx",
            type_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            auteur=self.admin,
        )
        piece = PieceEcrite.objects.create(
            projet=projet,
            modele=self.modele,
            intitule="Pièce avec GED",
            redacteur=self.admin,
            document_ged=document,
        )

        reponse = self.client.post(f"/api/pieces-ecrites/{piece.id}/session-bureautique/", {})

        self.assertEqual(reponse.status_code, status.HTTP_200_OK, reponse.data)
        self.assertIn("WOPISrc=", reponse.data["url_editeur"])
        self.assertIn(str(document.id), reponse.data["url_editeur"])
        self.assertNotIn("access_token=", reponse.data["url_editeur"])

    def test_super_admin_peut_supprimer_physiquement_une_piece_ecrite(self):
        projet = Projet.objects.create(
            reference="2026-0500",
            intitule="Projet pièce à supprimer",
            organisation=self.organisation,
            responsable=self.admin,
        )
        piece = PieceEcrite.objects.create(
            projet=projet,
            modele=self.modele,
            intitule="Pièce à supprimer",
            redacteur=self.admin,
        )

        reponse = self.client.delete(f"/api/pieces-ecrites/{piece.id}/")

        self.assertEqual(reponse.status_code, status.HTTP_200_OK, reponse.data)
        self.assertFalse(PieceEcrite.objects.filter(pk=piece.pk).exists())


class ApiPiecesEcritesCreationTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.organisation = Organisation.objects.create(
            code="ORG-PE",
            nom="Organisation pièces écrites",
            type_organisation="bureau_etudes",
        )
        self.utilisateur = Utilisateur.objects.create_user(
            courriel="pieces-ecrites@example.com",
            password="motdepasse123",
            prenom="Paul",
            nom="Rédacteur",
            organisation=self.organisation,
        )
        self.client.force_authenticate(self.utilisateur)
        self.projet = Projet.objects.create(
            reference="2026-0600",
            intitule="Projet pièce libre",
            organisation=self.organisation,
            responsable=self.utilisateur,
            cree_par=self.utilisateur,
        )

    def test_creation_piece_sans_modele_utilise_un_modele_libre(self):
        reponse = self.client.post(
            "/api/pieces-ecrites/",
            {
                "projet": str(self.projet.id),
                "intitule": "Note libre de cadrage",
            },
            format="json",
        )

        self.assertEqual(reponse.status_code, status.HTTP_201_CREATED, reponse.data)
        piece = PieceEcrite.objects.get(pk=reponse.data["id"])
        self.assertIsNotNone(piece.modele_id)
        self.assertEqual(piece.modele.code, "piece-libre-autre")
        self.assertEqual(piece.modele.type_document, "autre")
