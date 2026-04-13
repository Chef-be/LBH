"""Services métier pour la rédaction, la génération et l'export des pièces écrites."""

from __future__ import annotations

import re
import subprocess
import tempfile
import math
from io import BytesIO
from pathlib import Path
from typing import Iterable

from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.db.models import Q
from django.utils import timezone
from django.utils.text import slugify
from docx import Document as DocumentWord
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from lxml import etree, html as html_parser
from openpyxl import Workbook, load_workbook

from .models import PieceEcrite
from .office import assurer_gabarit_bureautique, extension_gabarit_modele


def televerser_image_editeur(fichier) -> str:
    horodatage = timezone.now()
    nom_base = Path(getattr(fichier, "name", "image")).stem[:80]
    extension = Path(getattr(fichier, "name", "")).suffix.lower() or ".png"
    chemin = (
        f"pieces_ecrites/editeur/{horodatage:%Y/%m}/"
        f"{slugify(nom_base) or 'image'}-{horodatage:%Y%m%d%H%M%S%f}{extension}"
    )
    chemin_stocke = default_storage.save(chemin, fichier)
    return default_storage.url(chemin_stocke)


def _normaliser_html_importe(html: str) -> str:
    fragments = html_parser.fragments_fromstring(html or "")
    sorties: list[str] = []

    for fragment in fragments:
        if isinstance(fragment, str):
            texte = fragment.strip()
            if texte:
                sorties.append(f"<p>{texte}</p>")
            continue

        for commentaire in fragment.xpath("//comment()"):
            parent = commentaire.getparent()
            if parent is not None:
                parent.remove(commentaire)

        for element in fragment.xpath(".//*"):
            etiquette = element.tag.lower() if isinstance(element.tag, str) else ""
            if etiquette in {"script", "style", "meta", "link", "xml"}:
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
                continue

            if etiquette in {"o:p", "v:shape", "v:imagedata"}:
                parent = element.getparent()
                if parent is not None:
                    index = parent.index(element)
                    for enfant in list(element):
                        parent.insert(index, enfant)
                        index += 1
                    parent.remove(element)
                continue

            classes = element.attrib.get("class", "")
            if classes:
                classes_filtrees = [
                    classe for classe in classes.split()
                    if not classe.lower().startswith("mso")
                ]
                if classes_filtrees:
                    element.attrib["class"] = " ".join(classes_filtrees)
                else:
                    element.attrib.pop("class", None)

            for attribut in list(element.attrib):
                if attribut.lower().startswith("xmlns") or attribut.lower().startswith("lang"):
                    element.attrib.pop(attribut, None)

        sorties.append(html_parser.tostring(fragment, encoding="unicode", method="html"))

    return "".join(sorties).strip()


def importer_fichier_word_en_html(fichier, construire_url_absolue=None) -> dict[str, object]:
    try:
        import mammoth
    except ImportError as exc:
        raise RuntimeError("La conversion Word n'est pas disponible sur ce serveur.") from exc

    def convertir_image(image):
        with image.open() as flux_image:
            contenu = flux_image.read()

        extension = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/webp": ".webp",
        }.get(image.content_type, ".png")

        fichier_image = ContentFile(contenu, name=f"image-importee{extension}")
        url = televerser_image_editeur(fichier_image)
        if construire_url_absolue:
            url = construire_url_absolue(url)
        return {"src": url}

    resultat = mammoth.convert_to_html(
        fichier,
        convert_image=mammoth.images.img_element(convertir_image),
    )
    html_converti = _normaliser_html_importe(resultat.value)
    return {
        "html": html_converti,
        "messages": [
            {"type": message.type, "message": message.message}
            for message in resultat.messages
        ],
    }


def _html_vers_texte(html: str) -> str:
    fragments = html_parser.fragments_fromstring(html or "")
    morceaux: list[str] = []
    for fragment in fragments:
        if isinstance(fragment, str):
            texte = fragment.strip()
        else:
            texte = " ".join(fragment.itertext()).strip()
        if texte:
            morceaux.append(texte)
    return "\n\n".join(morceaux)


def _balise_paragraphe_si_necessaire(contenu: str) -> str:
    texte = (contenu or "").strip()
    if not texte:
        return ""
    if "<" in texte and ">" in texte:
        return texte
    lignes = [ligne.strip() for ligne in texte.splitlines() if ligne.strip()]
    return "".join(f"<p>{ligne}</p>" for ligne in lignes)


def _texte_liste(valeur: Iterable[str] | str) -> str:
    if isinstance(valeur, str):
        return valeur
    return ", ".join(str(item) for item in valeur if item)


def _valeur_fusion_piece(piece: PieceEcrite, nom: str) -> str:
    if nom in piece.variables_personnalisees:
        valeur = piece.variables_personnalisees.get(nom)
        return "" if valeur is None else str(valeur)

    projet = piece.projet
    lot = piece.lot
    utilisateur = piece.redacteur
    valeurs = {
        "nom_projet": projet.intitule,
        "reference_projet": projet.reference,
        "description_projet": projet.description or "",
        "commune_projet": projet.commune or "",
        "departement_projet": projet.departement or "",
        "phase_projet": projet.get_phase_actuelle_display() if projet.phase_actuelle else "",
        "statut_projet": projet.get_statut_display(),
        "maitre_ouvrage": projet.maitre_ouvrage.nom if projet.maitre_ouvrage else "",
        "maitre_oeuvre": projet.maitre_oeuvre.nom if projet.maitre_oeuvre else "",
        "organisation": projet.organisation.nom if projet.organisation else "",
        "responsable_projet": projet.responsable.nom_complet if projet.responsable else "",
        "date_generation": timezone.localtime().strftime("%d/%m/%Y"),
        "lot_intitule": lot.intitule if lot else "",
        "lot_numero": str(lot.code) if lot else "",
        "redacteur_nom": utilisateur.nom_complet if utilisateur else "",
        "piece_intitule": piece.intitule,
        "modele_libelle": piece.modele.libelle if piece.modele_id else "",
    }
    return str(valeurs.get(nom, ""))


def construire_donnees_fusion_piece(piece: PieceEcrite) -> dict[str, str]:
    variables = piece.modele.variables_fusion or []
    donnees: dict[str, str] = {}

    for variable in variables:
        if not isinstance(variable, dict):
            continue
        nom = str(variable.get("nom") or "").strip()
        if not nom:
            continue
        donnees[nom] = _valeur_fusion_piece(piece, nom)

    for nom in (
        "nom_projet",
        "reference_projet",
        "commune_projet",
        "maitre_ouvrage",
        "maitre_oeuvre",
        "organisation",
        "responsable_projet",
        "date_generation",
        "lot_intitule",
        "lot_numero",
        "redacteur_nom",
        "piece_intitule",
        "modele_libelle",
        "contenu_principal",
    ):
        donnees.setdefault(nom, _valeur_fusion_piece(piece, nom))

    donnees["contenu_principal"] = _html_vers_texte(piece.contenu_html or generer_contenu_piece_depuis_articles(piece))

    return donnees


def _echapper_html(valeur: str) -> str:
    return (
        str(valeur)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def interpoler_variables_modele(texte: str, donnees_fusion: dict[str, str]) -> str:
    def remplacer(correspondance: re.Match[str]) -> str:
        nom = correspondance.group(1).strip()
        return _echapper_html(donnees_fusion.get(nom, ""))

    return re.sub(r"\{([^{}]+)\}", remplacer, texte or "")


def generer_contenu_piece_depuis_modele(piece: PieceEcrite) -> str:
    donnees_fusion = construire_donnees_fusion_piece(piece)
    contenu_modele = (piece.modele.contenu_modele_html or "").strip()
    if contenu_modele:
        return interpoler_variables_modele(contenu_modele, donnees_fusion)

    if piece.modele.type_document == "cctp":
        return _generer_cctp_metier(piece)
    if piece.modele.type_document == "lettre_candidature":
        return _generer_lettre_candidature(piece)
    if piece.modele.type_document == "memoire_technique":
        return _generer_memoire_technique(piece)
    if piece.modele.type_document == "planning_taches":
        return _generer_planning_taches(piece)
    if piece.modele.type_document == "rapport_analyse":
        return _generer_rapport_analyse(piece)

    variables = [v for v in (piece.modele.variables_fusion or []) if isinstance(v, dict) and v.get("nom")]
    sections = [
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        f"<p><strong>Modèle :</strong> {_echapper_html(piece.modele.libelle)}</p>",
        f"<p><strong>Projet :</strong> {_echapper_html(piece.projet.reference)} — {_echapper_html(piece.projet.intitule)}</p>",
    ]

    if piece.modele.description:
        sections.append(f"<p>{_echapper_html(interpoler_variables_modele(piece.modele.description, donnees_fusion))}</p>")

    if variables:
        sections.append("<h2>Données de fusion</h2>")
        sections.append("<table><tbody>")
        for variable in variables:
            nom = str(variable.get("nom") or "")
            libelle = str(variable.get("description") or nom)
            valeur = donnees_fusion.get(nom, "") or str(variable.get("exemple") or "")
            sections.append(
                "<tr>"
                f"<td><strong>{_echapper_html(libelle)}</strong></td>"
                f"<td>{_echapper_html(valeur)}</td>"
                "</tr>"
            )
        sections.append("</tbody></table>")

    if piece.modele.type_document == "cctp":
        sections.append("<h2>Prescriptions générales</h2>")
        sections.append(
            "<p>Les prestations seront exécutées conformément aux pièces du marché, "
            "aux règles de l’art et aux normes applicables au présent projet.</p>"
        )
    elif piece.modele.type_document in {"dpgf", "bpu", "dqe"}:
        sections.append("<h2>Cadre quantitatif et financier</h2>")
        sections.append(
            "<p>Les quantités, unités et hypothèses financières seront renseignées en cohérence "
            "avec le périmètre du projet et les découpages de lots.</p>"
        )
    elif piece.modele.type_document in {"rc", "ae", "ccap"}:
        sections.append("<h2>Cadre contractuel</h2>")
        sections.append(
            "<p>Le présent document constitue une base de rédaction contractuelle à compléter "
            "selon la consultation et le maître d’ouvrage concerné.</p>"
        )
    else:
        sections.append("<h2>Contenu initial</h2>")
        sections.append("<p>Ce document a été généré à partir du modèle sélectionné et des variables métier disponibles.</p>")

    return "".join(sections)


def _etude_economique_reference(piece: PieceEcrite):
    from applications.economie.models import EtudeEconomique

    etudes = EtudeEconomique.objects.filter(projet=piece.projet).prefetch_related("lignes")
    if piece.lot_id:
        etudes = etudes.filter(Q(lot_id=piece.lot_id) | Q(lot__isnull=True))
    return etudes.order_by("-version", "-date_modification").first()


def _lignes_dpgf(piece: PieceEcrite):
    etude = _etude_economique_reference(piece)
    if not etude:
        return []
    return list(etude.lignes.order_by("numero_ordre"))


def _contexte_client(piece: PieceEcrite) -> str:
    maitre_ouvrage = piece.projet.maitre_ouvrage
    if maitre_ouvrage and maitre_ouvrage.type_organisation == "maitre_ouvrage":
        return "maître d’ouvrage"
    if maitre_ouvrage and maitre_ouvrage.type_organisation == "entreprise":
        return "entreprise privée"
    return "client"


def _strategie_reponse_mission(piece: PieceEcrite) -> dict[str, object]:
    type_client = getattr(piece.projet.maitre_ouvrage, "type_organisation", "") if piece.projet.maitre_ouvrage else ""
    phase = piece.projet.phase_actuelle or ""

    if type_client == "maitre_ouvrage":
        return {
            "angle": "commande publique",
            "message": "La réponse doit démontrer la compréhension du besoin, la rigueur méthodologique, la fiabilité des moyens et la traçabilité des livrables.",
            "criteres": [
                ("Compréhension du besoin", 30),
                ("Méthodologie et organisation", 25),
                ("Moyens humains dédiés", 20),
                ("Maîtrise du planning", 15),
                ("Qualité, reporting et engagements", 10),
            ],
        }
    if type_client in {"entreprise", "partenaire"}:
        return {
            "angle": "mission privée / PME",
            "message": "La réponse doit rassurer sur la réactivité, la clarté du chiffrage, la maîtrise budgétaire et la capacité d'accompagnement décisionnel.",
            "criteres": [
                ("Réactivité et disponibilité", 25),
                ("Clarté des livrables", 20),
                ("Maîtrise du budget", 25),
                ("Souplesse d'intervention", 15),
                ("Coordination opérationnelle", 15),
            ],
        }
    if type_client == "sous_traitant" or phase == "exe":
        return {
            "angle": "co-traitance / sous-traitance / exécution",
            "message": "La réponse doit démontrer la fiabilité d'exécution, la coordination avec les autres intervenants et la capacité à tenir les interfaces calendaires.",
            "criteres": [
                ("Coordination avec les intervenants", 30),
                ("Fiabilité technique et documentaire", 25),
                ("Tenue des délais", 20),
                ("Gestion des interfaces", 15),
                ("Reporting et suivi", 10),
            ],
        }
    return {
        "angle": "mission générale",
        "message": "La réponse doit exposer une organisation lisible, des moyens adaptés et une capacité à sécuriser les coûts, délais et livrables.",
        "criteres": [
            ("Compréhension de la mission", 25),
            ("Organisation proposée", 25),
            ("Moyens affectés", 20),
            ("Sécurisation coûts et délais", 20),
            ("Qualité des livrables", 10),
        ],
    }


def _tableau_criteres_html(criteres: list[tuple[str, int]]) -> str:
    lignes = "".join(
        f"<tr><td>{_echapper_html(libelle)}</td><td>{poids} %</td></tr>"
        for libelle, poids in criteres
    )
    return (
        "<table><thead><tr><th>Critère valorisé</th><th>Pondération indicative</th></tr></thead>"
        f"<tbody>{lignes}</tbody></table>"
    )


def proposer_article_cctp_assiste(piece: PieceEcrite, donnees: dict[str, object]) -> dict[str, object]:
    from applications.bibliotheque.models import LignePrixBibliotheque

    ligne_reference = None
    ligne_reference_id = donnees.get("ligne_prix_reference")
    if ligne_reference_id:
        ligne_reference = LignePrixBibliotheque.objects.filter(pk=ligne_reference_id).first()

    chapitre = str(donnees.get("chapitre") or "").strip()
    numero_article = str(donnees.get("numero_article") or "").strip()
    intitule = str(donnees.get("intitule") or "").strip()
    niveau_detail = str(donnees.get("niveau_detail") or "standard").strip() or "standard"
    inclure_mise_en_oeuvre = bool(donnees.get("inclure_mise_en_oeuvre", True))
    inclure_controles = bool(donnees.get("inclure_controles", True))
    inclure_dechets = bool(donnees.get("inclure_dechets", True))
    prescriptions_complementaires = str(donnees.get("prescriptions_complementaires") or "").strip()

    if ligne_reference and not intitule:
        intitule = ligne_reference.designation_courte

    normes = list(getattr(ligne_reference, "normes_applicables", []) or [])
    tags = list(dict.fromkeys(
        valeur for valeur in [
            getattr(piece.lot, "intitule", "") if piece.lot_id else "",
            getattr(ligne_reference, "famille", "") if ligne_reference else "",
            getattr(ligne_reference, "lot", "") if ligne_reference else "",
        ]
        if valeur
    ))

    paragraphes: list[str] = []
    paragraphes.append(
        "<p>Le présent article définit la consistance des ouvrages, les conditions de fourniture, "
        "de mise en œuvre et de contrôle nécessaires à la parfaite exécution de la prestation, "
        "y compris toutes sujétions, accessoires et prestations induites.</p>"
    )

    if ligne_reference and ligne_reference.prescriptions_techniques:
        paragraphes.append("<h3>Consistance des prestations</h3>")
        paragraphes.append(_balise_paragraphe_si_necessaire(ligne_reference.prescriptions_techniques))
    elif intitule:
        paragraphes.append("<h3>Consistance des prestations</h3>")
        paragraphes.append(
            f"<p>La prestation « {_echapper_html(intitule)} » comprend l'ensemble des fournitures, "
            "fixations, accessoires, réglages, raccords, protections et sujétions nécessaires à un ouvrage fini.</p>"
        )

    if inclure_mise_en_oeuvre:
        paragraphes.append("<h3>Mise en œuvre</h3>")
        if ligne_reference and ligne_reference.cahier_des_charges_structure:
            items = []
            for bloc in ligne_reference.cahier_des_charges_structure:
                if not isinstance(bloc, dict):
                    continue
                titre = str(bloc.get("titre") or "").strip()
                contenu = str(bloc.get("contenu") or "").strip()
                if not contenu:
                    continue
                items.append(
                    f"<li><strong>{_echapper_html(titre or 'Prescription')} :</strong> {_echapper_html(contenu)}</li>"
                )
            if items:
                paragraphes.append("<ul>" + "".join(items) + "</ul>")
        elif ligne_reference and ligne_reference.phases_execution:
            paragraphes.append(
                "<ul>"
                + "".join(f"<li>{_echapper_html(str(phase))}</li>" for phase in ligne_reference.phases_execution if phase)
                + "</ul>"
            )
        else:
            paragraphes.append(
                "<p>La mise en œuvre devra respecter les prescriptions des fabricants, les règles de l'art, "
                "les documents techniques applicables et les exigences de coordination entre lots.</p>"
            )

    if inclure_controles:
        paragraphes.append("<h3>Contrôles, essais et documents à remettre</h3>")
        texte_controle = (
            "<p>L'entreprise remettra les fiches techniques, procès-verbaux d'essais, attestations de conformité, "
            "notices d'entretien, DOE et tout document de traçabilité nécessaire à la réception des ouvrages.</p>"
        )
        if niveau_detail == "detaille":
            texte_controle = (
                "<p>L'entreprise remettra avant exécution les fiches techniques, échantillons, plans d'atelier si nécessaire, "
                "puis en cours et fin d'exécution les procès-verbaux d'essais, constats de conformité, DOE, plans de récolement "
                "et notices d'entretien requis pour l'acceptation des ouvrages.</p>"
            )
        paragraphes.append(texte_controle)

    if inclure_dechets:
        paragraphes.append("<h3>Gestion des déchets et sujétions environnementales</h3>")
        dechets = list(getattr(ligne_reference, "dechets_generes", []) or []) if ligne_reference else []
        if dechets:
            paragraphes.append(
                "<ul>"
                + "".join(
                    f"<li>{_echapper_html(str(dechet))}</li>"
                    for dechet in dechets
                    if dechet
                )
                + "</ul>"
            )
        else:
            paragraphes.append(
                "<p>Le titulaire assurera le tri, l'évacuation réglementaire et la traçabilité des déchets issus de sa prestation, "
                "en cohérence avec le SOSED, le PGC ou tout dispositif environnemental applicable au marché.</p>"
            )

    if ligne_reference and ligne_reference.criteres_metre:
        paragraphes.append("<h3>Mode de métré et conditions de règlement</h3>")
        paragraphes.append(_balise_paragraphe_si_necessaire(ligne_reference.criteres_metre))

    if prescriptions_complementaires:
        paragraphes.append("<h3>Prescriptions complémentaires</h3>")
        paragraphes.append(_balise_paragraphe_si_necessaire(prescriptions_complementaires))

    if niveau_detail == "detaille":
        paragraphes.append(
            "<p><em>Les prestations non explicitement mentionnées mais nécessaires à la parfaite finition, à la sécurité, "
            "à la coordination ou à la conformité réglementaire sont réputées incluses dans le prix.</em></p>"
        )

    code_reference = ""
    if ligne_reference:
        suffixe = str(ligne_reference.code or ligne_reference.id)[:24].upper()
        code_reference = f"LBH-CCTP-{suffixe}"

    return {
        "piece_ecrite": str(piece.id),
        "chapitre": chapitre,
        "numero_article": numero_article,
        "code_reference": code_reference,
        "intitule": intitule or "Article technique",
        "corps_article": "".join(paragraphes),
        "source": getattr(ligne_reference, "source", "") if ligne_reference else "Proposition CCTP LBH",
        "source_url": getattr(ligne_reference, "url_source", "") if ligne_reference else "",
        "ligne_prix_reference": str(ligne_reference.id) if ligne_reference else None,
        "normes_applicables": normes,
        "tags": tags,
        "est_dans_bibliotheque": bool(donnees.get("est_dans_bibliotheque", False)),
    }


def _generer_cctp_metier(piece: PieceEcrite) -> str:
    articles = piece.articles.order_by("chapitre", "numero_article", "date_creation")
    lignes = _lignes_dpgf(piece)
    maitre_ouvrage = _valeur_fusion_piece(piece, "maitre_ouvrage")
    sections = [
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        "<p><strong>Objet du lot :</strong> prescriptions techniques, modalités d’exécution, qualité attendue et limites de prestations.</p>",
        f"<p><strong>Projet :</strong> {_echapper_html(piece.projet.reference)} — {_echapper_html(piece.projet.intitule)}</p>",
    ]
    if piece.projet.description:
        sections.append(f"<p>{_echapper_html(piece.projet.description)}</p>")
    if maitre_ouvrage:
        sections.append(f"<p><strong>Maître d’ouvrage :</strong> {_echapper_html(maitre_ouvrage)}</p>")

    sections.append("<h2>Dispositions générales</h2>")
    sections.append(
        "<p>Les ouvrages devront être exécutés conformément aux pièces du marché, aux règles de l’art, "
        "aux normes en vigueur et aux contraintes propres au site. L’entreprise devra remettre tous les "
        "documents d’exécution, fiches techniques, plans d’atelier et notices nécessaires à la bonne compréhension du lot.</p>"
    )

    if articles.exists():
        sections.append("<h2>Articles techniques</h2>")
        sections.append(generer_contenu_piece_depuis_articles(piece))
    elif lignes:
        sections.append("<h2>Prescriptions par poste quantitatif</h2>")
        for ligne in lignes:
            sections.append(f"<h3>{_echapper_html(ligne.code or str(ligne.numero_ordre))} — {_echapper_html(ligne.designation)}</h3>")
            sections.append(
                "<p>Le titulaire devra exécuter ce poste avec tous accessoires, sujétions de transport, "
                "de mise en œuvre, de protection, de contrôle et de réception nécessaires à un ouvrage fini.</p>"
            )
            if ligne.observations:
                sections.append(f"<p><em>Observations économiques et techniques :</em> {_echapper_html(ligne.observations)}</p>")
    else:
        sections.append("<p>Compléter les articles techniques du lot avant diffusion du CCTP.</p>")

    sections.append("<h2>Contrôles, essais et réception</h2>")
    sections.append(
        "<p>L’entreprise devra fournir les procès-verbaux d’essais, les fiches de contrôle, les DOE, "
        "ainsi que tout justificatif de conformité nécessaire à la réception des ouvrages.</p>"
    )
    return "".join(sections)


def _generer_lettre_candidature(piece: PieceEcrite) -> str:
    organisation = piece.projet.organisation
    maitre_ouvrage = piece.projet.maitre_ouvrage
    strategie = _strategie_reponse_mission(piece)
    return "".join([
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        "<p>Objet : candidature à la consultation</p>",
        "<p>Madame, Monsieur,</p>",
        f"<p>Par la présente, {_echapper_html(organisation.nom if organisation else 'notre bureau d’études')} "
        f"vous confirme sa candidature pour l’opération <strong>{_echapper_html(piece.projet.intitule)}</strong> "
        f"référencée <strong>{_echapper_html(piece.projet.reference)}</strong>.</p>",
        f"<p>Notre intervention porte sur une mission adaptée au contexte {_echapper_html(_contexte_client(piece))}, "
        "avec mobilisation de compétences en économie de la construction, rédaction des pièces écrites, "
        "analyse d’offres, suivi contractuel et accompagnement technique suivant la phase du projet.</p>",
        f"<p>Le dossier de candidature vise à démontrer nos capacités techniques, nos références comparables, "
        f"notre organisation de mission ainsi que notre compréhension des attentes du maître d’ouvrage "
        f"{_echapper_html(maitre_ouvrage.nom if maitre_ouvrage else '')}.</p>",
        f"<p><strong>Angle de réponse privilégié :</strong> {_echapper_html(strategie['angle'])}.</p>",
        f"<p>{_echapper_html(strategie['message'])}</p>",
        "<h2>Critères à valoriser dans la candidature</h2>",
        _tableau_criteres_html(strategie["criteres"]),
        "<p>Nous restons à votre disposition pour toute précision complémentaire et vous prions d’agréer, Madame, Monsieur, l’expression de nos salutations distinguées.</p>",
    ])


def _generer_memoire_technique(piece: PieceEcrite) -> str:
    lignes = _lignes_dpgf(piece)
    strategie = _strategie_reponse_mission(piece)
    sections = [
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        "<h2>Compréhension du besoin</h2>",
        f"<p>Le projet {_echapper_html(piece.projet.reference)} concerne {_echapper_html(piece.projet.intitule)}. "
        "Le mémoire technique expose notre compréhension de l’opération, notre méthode, notre organisation et nos engagements.</p>",
        "<h2>Positionnement de la réponse</h2>",
        f"<p><strong>Contexte ciblé :</strong> {_echapper_html(strategie['angle'])}. { _echapper_html(strategie['message']) }</p>",
        _tableau_criteres_html(strategie["criteres"]),
        "<h2>Organisation de mission</h2>",
        "<p>La mission sera pilotée par un référent unique, assisté des compétences d’économie, rédaction, analyse documentaire et suivi des interfaces. "
        "Les livrables feront l’objet d’une vérification croisée avant transmission.</p>",
        "<h2>Méthodologie d’exécution</h2>",
        "<ul>"
        "<li>Analyse préalable du dossier, des contraintes de site et du programme.</li>"
        "<li>Production et contrôle des pièces écrites, quantitatifs et chiffrages.</li>"
        "<li>Gestion des interfaces techniques, administratives et calendaires.</li>"
        "<li>Traçabilité des hypothèses, variantes et arbitrages.</li>"
        "</ul>",
        "<h2>Moyens affectés</h2>",
        "<p>Les moyens humains et numériques seront calibrés selon la complexité du projet, la phase de mission et la stratégie de consultation retenue.</p>",
    ]
    if lignes:
        sections.append("<h2>Postes sensibles identifiés</h2>")
        sections.append("<ul>")
        for ligne in lignes[:12]:
            sections.append(f"<li>{_echapper_html(ligne.designation)}</li>")
        sections.append("</ul>")
    sections.append("<h2>Qualité, délais et livrables</h2>")
    sections.append(
        "<p>Les livrables seront diffusés dans un format contrôlé, versionné et compatible avec les exigences du maître d’ouvrage. "
        "Chaque production intégrera une vérification de cohérence technique, économique et contractuelle.</p>"
    )
    return "".join(sections)


def _duree_tache_jours(ligne) -> int:
    quantite = getattr(ligne, "quantite_prevue", None) or getattr(ligne, "quantite", None)
    if ligne.temps_main_oeuvre and quantite:
        heures = float(ligne.temps_main_oeuvre) * float(quantite)
        return max(1, math.ceil(heures / 7))
    if quantite:
        return max(1, math.ceil(float(quantite) / 10))
    return 1


def _generer_planning_taches(piece: PieceEcrite) -> str:
    lignes = _lignes_dpgf(piece)
    sections = [
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        "<p>Planning prévisionnel des tâches construit à partir des lignes économiques disponibles.</p>",
        "<table><thead><tr><th>Ordre</th><th>Tâche</th><th>Base</th><th>Durée estimative</th></tr></thead><tbody>",
    ]
    if lignes:
        for ligne in lignes:
            sections.append(
                f"<tr><td>{ligne.numero_ordre}</td>"
                f"<td>{_echapper_html(ligne.designation)}</td>"
                f"<td>{_echapper_html(ligne.unite)} · {_echapper_html(str(ligne.quantite_prevue))}</td>"
                f"<td>{_duree_tache_jours(ligne)} jour(s)</td></tr>"
            )
    else:
        sections.append("<tr><td colspan='4'>Aucune ligne DPGF / étude économique exploitable pour générer le planning.</td></tr>")
    sections.append("</tbody></table>")
    sections.append(
        "<p><em>Ce planning constitue une base de phasage. Il doit être ajusté selon les interfaces chantier, "
        "les approvisionnements, les sujétions de site et les contraintes de coactivité.</em></p>"
    )
    return "".join(sections)


def _generer_rapport_analyse(piece: PieceEcrite) -> str:
    from applications.documents.models import Document

    lignes = _lignes_dpgf(piece)
    documents = Document.objects.filter(projet=piece.projet).order_by("-date_modification")[:10]
    strategie = _strategie_reponse_mission(piece)
    sections = [
        f"<h1>{_echapper_html(piece.intitule)}</h1>",
        "<h2>Contexte de l'analyse</h2>",
        f"<p>Le présent rapport synthétise les éléments significatifs identifiés sur le projet "
        f"{_echapper_html(piece.projet.reference)} — {_echapper_html(piece.projet.intitule)}.</p>",
        "<h2>Grille d'analyse prioritaire</h2>",
        f"<p>{_echapper_html(strategie['message'])}</p>",
        _tableau_criteres_html(strategie["criteres"]),
    ]
    if lignes:
        sections.append("<h2>Analyse économique</h2><ul>")
        for ligne in lignes[:10]:
            sections.append(
                f"<li>{_echapper_html(ligne.designation)} : "
                f"quantité {_echapper_html(str(ligne.quantite_prevue))} {_echapper_html(ligne.unite)}, "
                f"prix de vente unitaire {_echapper_html(str(ligne.prix_vente_unitaire))} €.</li>"
            )
        sections.append("</ul>")
    if documents:
        sections.append("<h2>Documents examinés</h2><ul>")
        for document in documents:
            sections.append(f"<li>{_echapper_html(document.reference)} — {_echapper_html(document.intitule)}</li>")
        sections.append("</ul>")
    sections.append("<h2>Points d'attention et recommandations</h2>")
    sections.append(
        "<p>Les éléments présentant un impact potentiel sur les coûts, délais, interfaces ou conformité réglementaire "
        "doivent faire l’objet d’un suivi prioritaire avec traçabilité des décisions, justificatifs et versions de documents.</p>"
    )
    return "".join(sections)


def generer_contenu_piece_depuis_articles(piece: PieceEcrite) -> str:
    articles = piece.articles.order_by("chapitre", "numero_article", "date_creation")
    blocs = [
        f"<h1>{piece.intitule}</h1>",
        "<section>",
        f"<p><strong>Projet :</strong> {piece.projet.reference} — {piece.projet.intitule}</p>",
    ]
    if piece.modele_id:
        blocs.append(f"<p><strong>Modèle :</strong> {piece.modele.libelle}</p>")
    blocs.append("</section>")

    for article in articles:
        titre = " — ".join(
            partie for partie in [
                ".".join(partie for partie in [article.chapitre, article.numero_article] if partie),
                article.intitule,
            ] if partie
        )
        blocs.append("<section>")
        blocs.append(f"<h2>{titre or 'Article'}</h2>")
        corps = _balise_paragraphe_si_necessaire(article.corps_article)
        if corps:
            blocs.append(corps)
        if article.normes_applicables:
            blocs.append(f"<p><em>Normes applicables : {_texte_liste(article.normes_applicables)}</em></p>")
        blocs.append("</section>")

    if not articles.exists():
        blocs.append("<p>Aucun article n'est encore renseigné.</p>")
    return "".join(blocs)


def regenerer_piece_ecrite(piece: PieceEcrite) -> PieceEcrite:
    piece.contenu_html = generer_contenu_piece_depuis_articles(piece)
    piece.date_generation = timezone.now()
    piece.save(update_fields=["contenu_html", "date_generation", "date_modification"])
    return piece


def generer_piece_depuis_modele(piece: PieceEcrite) -> PieceEcrite:
    piece.contenu_html = generer_contenu_piece_depuis_modele(piece)
    piece.date_generation = timezone.now()
    piece.save(update_fields=["contenu_html", "date_generation", "date_modification", "variables_personnalisees"])
    return piece


def _iterer_blocs_docx(document: DocumentWord):
    for paragraphe in document.paragraphs:
        yield paragraphe
    for tableau in document.tables:
        for ligne in tableau.rows:
            for cellule in ligne.cells:
                for paragraphe in cellule.paragraphs:
                    yield paragraphe
    for section in document.sections:
        for paragraphe in section.header.paragraphs:
            yield paragraphe
        for paragraphe in section.footer.paragraphs:
            yield paragraphe


def _remplacer_placeholders_docx(document: DocumentWord, donnees_fusion: dict[str, str]):
    remplacements = {f"{{{cle}}}": valeur for cle, valeur in donnees_fusion.items()}
    for paragraphe in _iterer_blocs_docx(document):
        if not paragraphe.text:
            continue
        remplacement_effectue = False
        for run in paragraphe.runs:
            texte_run = run.text or ""
            nouveau_run = texte_run
            for placeholder, valeur in remplacements.items():
                nouveau_run = nouveau_run.replace(placeholder, valeur)
            if nouveau_run != texte_run:
                run.text = nouveau_run
                remplacement_effectue = True

        if remplacement_effectue or "{" not in paragraphe.text:
            continue

        texte = paragraphe.text
        nouveau = texte
        for placeholder, valeur in remplacements.items():
            nouveau = nouveau.replace(placeholder, valeur)
        if nouveau == texte:
            continue
        if paragraphe.runs:
            premier = paragraphe.runs[0]
            premier.text = nouveau
            for run in paragraphe.runs[1:]:
                run.text = ""
        else:
            paragraphe.add_run(nouveau)


def _ajouter_runs(paragraphe, element, gras=False, italique=False, souligne=False, barre=False):
    if element.text:
        run = paragraphe.add_run(element.text)
        run.bold = gras
        run.italic = italique
        run.underline = souligne
        run.font.strike = barre

    for enfant in element:
        etiquette = enfant.tag.lower() if isinstance(enfant.tag, str) else ""
        enfant_gras = gras or etiquette in {"strong", "b"}
        enfant_italique = italique or etiquette in {"em", "i"}
        enfant_souligne = souligne or etiquette == "u"
        enfant_barre = barre or etiquette in {"s", "strike"}

        if etiquette == "br":
            paragraphe.add_run().add_break()
        elif etiquette == "a":
            run = paragraphe.add_run("".join(enfant.itertext()))
            run.bold = enfant_gras
            run.italic = enfant_italique
            run.underline = True
        elif etiquette in {"code"}:
            run = paragraphe.add_run("".join(enfant.itertext()))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            _ajouter_runs(paragraphe, enfant, enfant_gras, enfant_italique, enfant_souligne, enfant_barre)

        if enfant.tail:
            run = paragraphe.add_run(enfant.tail)
            run.bold = gras
            run.italic = italique
            run.underline = souligne
            run.font.strike = barre


def _ajouter_element_word(document: DocumentWord, element):
    etiquette = element.tag.lower() if isinstance(element.tag, str) else ""
    if etiquette in {"h1", "h2", "h3"}:
        niveau = {"h1": 1, "h2": 2, "h3": 3}[etiquette]
        paragraphe = document.add_heading(level=niveau)
        _ajouter_runs(paragraphe, element)
        return

    if etiquette == "p":
        paragraphe = document.add_paragraph()
        _ajouter_runs(paragraphe, element)
        return

    if etiquette == "blockquote":
        paragraphe = document.add_paragraph()
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraphe.paragraph_format.left_indent = Inches(0.3)
        _ajouter_runs(paragraphe, element, italique=True)
        return

    if etiquette == "pre":
        paragraphe = document.add_paragraph()
        run = paragraphe.add_run("".join(element.itertext()))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        return

    if etiquette in {"ul", "ol"}:
        style = "List Bullet" if etiquette == "ul" else "List Number"
        for item in element:
            if not isinstance(item.tag, str) or item.tag.lower() != "li":
                continue
            paragraphe = document.add_paragraph(style=style)
            _ajouter_runs(paragraphe, item)
        return

    if etiquette == "table":
        lignes = [enfant for enfant in element if isinstance(enfant.tag, str) and enfant.tag.lower() in {"tbody", "thead", "tr"}]
        lignes_exp = []
        for ligne in lignes:
            if ligne.tag.lower() in {"tbody", "thead"}:
                lignes_exp.extend([f for f in ligne if isinstance(f.tag, str) and f.tag.lower() == "tr"])
            else:
                lignes_exp.append(ligne)
        if lignes_exp:
            premiere = lignes_exp[0]
            colonnes = max(1, len([c for c in premiere if isinstance(c.tag, str)]))
            tableau = document.add_table(rows=0, cols=colonnes)
            tableau.style = "Table Grid"
            for ligne in lignes_exp:
                cellules = [c for c in ligne if isinstance(c.tag, str)]
                row = tableau.add_row().cells
                for index, cellule in enumerate(cellules[:colonnes]):
                    row[index].text = " ".join(cellule.itertext()).strip()
        return

    if etiquette == "hr":
        document.add_paragraph("------------------------------------------------------------")
        return

    if etiquette in {"section", "div", "body"}:
        for enfant in element:
            if isinstance(enfant.tag, str):
                _ajouter_element_word(document, enfant)
        return

    if etiquette == "li":
        paragraphe = document.add_paragraph(style="List Bullet")
        _ajouter_runs(paragraphe, element)
        return

    if isinstance(element, etree._Element):
        paragraphe = document.add_paragraph()
        _ajouter_runs(paragraphe, element)


def generer_docx_piece(piece: PieceEcrite) -> bytes:
    if piece.modele_id and piece.modele.gabarit and piece.modele.gabarit.name.lower().endswith(".docx"):
        with piece.modele.gabarit.open("rb") as fichier_modele:
            document = DocumentWord(BytesIO(fichier_modele.read()))
        _remplacer_placeholders_docx(document, construire_donnees_fusion_piece(piece))
        flux = BytesIO()
        document.save(flux)
        return flux.getvalue()

    document = DocumentWord()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    titre = document.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre.add_run(piece.intitule)
    run_titre.bold = True
    run_titre.font.size = Pt(18)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{piece.projet.reference} — {piece.projet.intitule}")

    fragments = html_parser.fragments_fromstring(piece.contenu_html or generer_contenu_piece_depuis_articles(piece))
    for fragment in fragments:
        if isinstance(fragment, str):
            texte = fragment.strip()
            if texte:
                document.add_paragraph(texte)
            continue
        _ajouter_element_word(document, fragment)

    flux = BytesIO()
    document.save(flux)
    return flux.getvalue()


def generer_xlsx_piece(piece: PieceEcrite) -> bytes:
    donnees_fusion = construire_donnees_fusion_piece(piece)

    if piece.modele_id:
        assurer_gabarit_bureautique(piece.modele)

    if piece.modele_id and piece.modele.gabarit and piece.modele.gabarit.name.lower().endswith(".xlsx"):
        with piece.modele.gabarit.open("rb") as fichier_modele:
            classeur = load_workbook(BytesIO(fichier_modele.read()))
    else:
        classeur = Workbook()
        feuille = classeur.active
        feuille.title = "Modele"
        feuille["A1"] = "Projet"
        feuille["B1"] = "{reference_projet} - {nom_projet}"
        feuille["A2"] = "Lot"
        feuille["B2"] = "{lot_numero} - {lot_intitule}"

    for feuille in classeur.worksheets:
        for ligne in feuille.iter_rows():
            for cellule in ligne:
                if isinstance(cellule.value, str):
                    nouveau = cellule.value
                    for cle, valeur in donnees_fusion.items():
                        nouveau = nouveau.replace(f"{{{cle}}}", valeur)
                    cellule.value = nouveau

    flux = BytesIO()
    classeur.save(flux)
    return flux.getvalue()


def generer_pdf_piece(piece: PieceEcrite) -> bytes:
    extension_source = ".docx"
    contenu_source = generer_docx_piece(piece)

    if piece.modele_id:
        extension_modele = extension_gabarit_modele(piece.modele)
        if extension_modele == ".xlsx":
            extension_source = ".xlsx"
            contenu_source = generer_xlsx_piece(piece)

    with tempfile.TemporaryDirectory(prefix="lbh-piece-") as dossier_temp:
        chemin_source = Path(dossier_temp) / f"piece{extension_source}"
        chemin_pdf = Path(dossier_temp) / "piece.pdf"
        chemin_source.write_bytes(contenu_source)

        commande = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            dossier_temp,
            str(chemin_source),
        ]
        resultat = subprocess.run(
            commande,
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if resultat.returncode != 0 or not chemin_pdf.exists():
            detail = (resultat.stderr or resultat.stdout or "").strip()
            raise RuntimeError(detail or "Échec de conversion PDF via LibreOffice.")

        return chemin_pdf.read_bytes()


def _nom_fichier_export(piece: PieceEcrite, extension: str) -> str:
    base = f"{piece.projet.reference}-{piece.intitule}".strip()
    base = "".join(car if car.isalnum() or car in {"-", "_"} else "-" for car in base)
    base = "-".join(segment for segment in base.split("-") if segment)
    return f"{base[:120] or 'piece-ecrite'}.{extension}"


def exporter_piece_ecrite(piece: PieceEcrite, format_sortie: str) -> tuple[bytes, str, str]:
    if format_sortie == "docx":
        contenu = generer_docx_piece(piece)
        type_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format_sortie == "xlsx":
        contenu = generer_xlsx_piece(piece)
        type_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif format_sortie == "pdf":
        contenu = generer_pdf_piece(piece)
        type_mime = "application/pdf"
    else:
        raise ValueError(f"Format d'export non supporté : {format_sortie}")

    nom_fichier = _nom_fichier_export(piece, format_sortie)
    piece.fichier_genere.save(
        str(Path("pieces_ecrites") / nom_fichier),
        ContentFile(contenu),
        save=False,
    )
    piece.date_generation = timezone.now()
    from applications.documents.services import enregistrer_document_genere_projet

    document_ged = enregistrer_document_genere_projet(
        projet=piece.projet,
        reference=f"{piece.projet.reference}-{piece.modele.type_document.upper()}",
        intitule=piece.intitule,
        type_document_code=piece.modele.type_document.upper(),
        type_document_libelle=piece.modele.get_type_document_display(),
        auteur=piece.redacteur,
        contenu=contenu,
        nom_fichier=nom_fichier,
        type_mime=type_mime,
        document_existant=piece.document_ged,
        statut="valide" if piece.statut in {"valide", "diffuse"} else "brouillon",
        contexte_generation=f"piece-ecrite:{piece.modele.type_document}",
        confidentiel=False,
    )
    piece.document_ged = document_ged
    piece.save(update_fields=["fichier_genere", "document_ged", "date_generation", "date_modification"])
    return contenu, type_mime, nom_fichier


def generer_cctp_depuis_bibliotheque(
    projet,
    intitule: str,
    lots_numeros: list[str],
    prescriptions_exclues_ids: list[str],
    variables: dict,
    utilisateur=None,
) -> PieceEcrite:
    """
    Génère un CCTP Word complet depuis la bibliothèque de prescriptions.
    Retourne une PieceEcrite avec le fichier Word généré dans MinIO.
    """
    from .models import LotCCTP, PrescriptionCCTP, ModeleDocument

    # Récupérer les lots dans l'ordre
    lots = LotCCTP.objects.filter(
        code__in=lots_numeros, est_actif=True
    ).prefetch_related(
        "chapitres__prescriptions"
    ).order_by("ordre", "code")

    # Construire le document Word
    document = DocumentWord()

    # Titre principal
    document.add_heading(intitule or "Cahier des Clauses Techniques Particulières", level=0)

    # En-tête du projet
    variables_completes = {
        "nom_projet": getattr(projet, "intitule", ""),
        "reference_projet": getattr(projet, "reference", ""),
        "maitre_ouvrage": getattr(getattr(projet, "organisation", None), "raison_sociale", ""),
        "date_generation": timezone.now().strftime("%d/%m/%Y"),
        **variables,
    }

    entete = document.add_paragraph()
    entete.add_run(f"Projet : {variables_completes.get('nom_projet', '')}").bold = True
    document.add_paragraph(f"Référence : {variables_completes.get('reference_projet', '')}")
    document.add_paragraph(f"Maître d'ouvrage : {variables_completes.get('maitre_ouvrage', '')}")
    document.add_paragraph(f"Date d'établissement : {variables_completes.get('date_generation', '')}")
    document.add_paragraph("")

    exclues_ids = set(str(i) for i in prescriptions_exclues_ids)

    for lot in lots:
        document.add_heading(f"LOT {lot.code} — {lot.intitule.upper()}", level=1)

        for chapitre in lot.chapitres.all().order_by("ordre"):
            document.add_heading(f"{chapitre.numero} — {chapitre.intitule}", level=2)

            for prescrip in chapitre.prescriptions.filter(est_actif=True).order_by("ordre"):
                if str(prescrip.pk) in exclues_ids:
                    continue

                # Titre de l'article
                document.add_heading(prescrip.intitule, level=3)

                # Corps avec remplacement de variables
                corps = prescrip.corps
                for var, valeur in variables_completes.items():
                    corps = corps.replace(f"{{{var}}}", str(valeur))
                # Remplacer variables non résolues par leur valeur par défaut
                corps = re.sub(r"\{([^}:]+):-([^}]*)\}", r"\2", corps)
                corps = re.sub(r"\{([^}]+)\}", r"[\1]", corps)

                document.add_paragraph(corps)

                # Normes
                if prescrip.normes:
                    p = document.add_paragraph()
                    p.add_run("Normes applicables : ").italic = True
                    p.add_run(", ".join(prescrip.normes)).italic = True

    # Sauvegarder le document Word en mémoire
    flux = BytesIO()
    document.save(flux)
    contenu_word = flux.getvalue()

    # Trouver ou créer un modèle CCTP générique
    modele, _ = ModeleDocument.objects.get_or_create(
        code="CCTP-GENERE",
        defaults={
            "libelle": "CCTP généré automatiquement",
            "type_document": "cctp",
            "description": "Modèle utilisé pour les CCTP générés depuis la bibliothèque de prescriptions.",
        }
    )

    # Créer la PieceEcrite
    piece = PieceEcrite.objects.create(
        projet=projet,
        modele=modele,
        intitule=intitule,
        statut="brouillon",
        redacteur=utilisateur,
    )
    nom_fichier = f"cctp-{slugify(intitule) or 'document'}-{timezone.now():%Y%m%d}.docx"
    piece.fichier_genere.save(
        f"pieces_ecrites/{timezone.now():%Y/%m}/{nom_fichier}",
        ContentFile(contenu_word),
        save=True,
    )
    piece.date_generation = timezone.now()
    piece.save()

    return piece


# ---------------------------------------------------------------------------
# A4 — Renumérotation automatique des articles CCTP (Widloecher & Cusant)
# ---------------------------------------------------------------------------

def renumeroter_articles_cctp(piece: PieceEcrite) -> int:
    """
    Renumérotation automatique des articles d'une pièce CCTP selon la structure
    Widloecher & Cusant (3e éd.) :
      Partie I — Dispositions générales : I.1, I.2, I.3…
      Partie II — Descriptif des ouvrages : II.{chapitre}.{rang} ou II.{rang}

    Les articles sont regroupés par chapitre (champ `chapitre`).
    Les chapitres vides ou « Dispositions générales » → Partie I.
    Les autres chapitres → Partie II numérotés dans l'ordre d'apparition.

    Retourne le nombre d'articles mis à jour.
    """
    from .models import ArticleCCTP

    articles = list(
        piece.articles.select_related("lot").order_by("chapitre", "numero_article", "date_creation")
    )
    if not articles:
        return 0

    # Séparer dispositions générales et descriptif
    CHAPITRES_GENERAUX = {
        "", "dispositions générales", "generalites", "généralités",
        "objet du marché", "documents de référence", "normes applicables",
        "prescriptions générales", "objet", "généralités et objet",
    }

    groupe_general: list[ArticleCCTP] = []
    chapitres: dict[str, list[ArticleCCTP]] = {}

    for article in articles:
        cle = (article.chapitre or "").strip().lower()
        if cle in CHAPITRES_GENERAUX:
            groupe_general.append(article)
        else:
            chapitres.setdefault(article.chapitre or "Descriptif", []).append(article)

    a_mettre_a_jour: list[ArticleCCTP] = []

    # Partie I — Dispositions générales
    for rang, article in enumerate(groupe_general, start=1):
        nouveau_numero = f"I.{rang}"
        if article.numero_article != nouveau_numero:
            article.numero_article = nouveau_numero
            a_mettre_a_jour.append(article)

    # Partie II — Descriptif des ouvrages
    for num_chap, (intitule_chapitre, liste_articles) in enumerate(chapitres.items(), start=1):
        for rang, article in enumerate(liste_articles, start=1):
            nouveau_numero = f"II.{num_chap}.{rang}"
            if article.numero_article != nouveau_numero:
                article.numero_article = nouveau_numero
                a_mettre_a_jour.append(article)

    if a_mettre_a_jour:
        ArticleCCTP.objects.bulk_update(a_mettre_a_jour, ["numero_article"])

    return len(a_mettre_a_jour)


# ---------------------------------------------------------------------------
# A5 — Export DPGF / BPU depuis les articles CCTP
# ---------------------------------------------------------------------------

def _styles_excel_dpgf(wb):
    """Retourne un dict de styles nommés pour le DPGF/BPU."""
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side, numbers

    fin_mince = Side(style="thin", color="C0C8D8")
    bordure = Border(left=fin_mince, right=fin_mince, top=fin_mince, bottom=fin_mince)

    return {
        "entete": {
            "font": Font(bold=True, color="FFFFFF", size=10),
            "fill": PatternFill("solid", fgColor="2D4A8A"),
            "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
            "border": bordure,
        },
        "lot": {
            "font": Font(bold=True, color="FFFFFF", size=9),
            "fill": PatternFill("solid", fgColor="4A6FA5"),
            "alignment": Alignment(horizontal="left", vertical="center"),
            "border": bordure,
        },
        "chapitre": {
            "font": Font(bold=True, size=9, color="1E3A5F"),
            "fill": PatternFill("solid", fgColor="D6E4F0"),
            "alignment": Alignment(horizontal="left", vertical="center"),
            "border": bordure,
        },
        "article": {
            "font": Font(size=9),
            "alignment": Alignment(horizontal="left", vertical="top", wrap_text=True),
            "border": bordure,
        },
        "numero": {
            "font": Font(size=9, bold=True, color="2D4A8A"),
            "alignment": Alignment(horizontal="center", vertical="top"),
            "border": bordure,
        },
        "montant": {
            "font": Font(size=9),
            "alignment": Alignment(horizontal="right", vertical="top"),
            "border": bordure,
            "number_format": "#,##0.00",
        },
        "total": {
            "font": Font(bold=True, size=9),
            "fill": PatternFill("solid", fgColor="EFF4FB"),
            "alignment": Alignment(horizontal="right", vertical="center"),
            "border": bordure,
            "number_format": "#,##0.00",
        },
    }


def _appliquer_style(cellule, style: dict):
    from openpyxl.styles import Font, PatternFill, Alignment, Border
    for attr, valeur in style.items():
        if attr == "number_format":
            cellule.number_format = valeur
        else:
            setattr(cellule, attr, valeur)


def exporter_articles_dpgf(piece: PieceEcrite) -> bytes:
    """
    Génère un fichier DPGF (.xlsx) depuis les articles CCTP d'une pièce écrite.
    Chaque article CCTP = une ligne DPGF avec : code, désignation, unité, quantité, PU, montant.
    Les articles sont regroupés par lot CCTP puis par chapitre.
    Compatible avec le BPU (désignation seule) et la DQE (avec quantités).
    """
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    articles = list(
        piece.articles.select_related("lot").order_by(
            "lot__ordre", "lot__code", "chapitre", "numero_article"
        )
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "DPGF"

    styles = _styles_excel_dpgf(wb)

    # Informations projet en en-tête
    projet = piece.projet
    ws.merge_cells("A1:F1")
    titre = ws["A1"]
    titre.value = f"DÉCOMPOSITION DU PRIX GLOBAL ET FORFAITAIRE — {piece.intitule.upper()}"
    _appliquer_style(titre, {
        "font": Font(bold=True, size=12, color="1E3A5F"),
        "alignment": Alignment(horizontal="center", vertical="center"),
    })
    ws.row_dimensions[1].height = 20

    ws.merge_cells("A2:F2")
    sous_titre = ws["A2"]
    lot_projet = piece.lot
    sous_titre.value = (
        f"Projet : {projet.reference} — {projet.intitule}"
        + (f"  |  Lot : {lot_projet.numero} {lot_projet.intitule}" if lot_projet else "")
        + f"  |  Date : {timezone.localtime().strftime('%d/%m/%Y')}"
    )
    _appliquer_style(sous_titre, {
        "font": Font(size=9, color="4A6FA5"),
        "alignment": Alignment(horizontal="center"),
    })
    ws.row_dimensions[2].height = 14

    # En-têtes colonnes (ligne 4)
    entetes = ["N°", "DÉSIGNATION DES OUVRAGES", "UNITÉ", "QUANTITÉ", "PRIX UNITAIRE HT (€)", "MONTANT HT (€)"]
    largeurs = [12, 55, 10, 12, 20, 18]
    for col, (libelle, largeur) in enumerate(zip(entetes, largeurs), start=1):
        cellule = ws.cell(row=4, column=col, value=libelle)
        _appliquer_style(cellule, styles["entete"])
        ws.column_dimensions[get_column_letter(col)].width = largeur
    ws.row_dimensions[4].height = 30

    ligne = 5
    lot_courant = None
    chapitre_courant = None
    premiere_ligne_lot = None
    sommes_lots: list[tuple[int, int]] = []  # (ligne_debut, ligne_fin)

    for article in articles:
        lot_cctp = article.lot
        lot_libelle = f"{lot_cctp.code} — {lot_cctp.intitule}" if lot_cctp else "Sans lot"

        # Séparateur de lot
        if lot_libelle != lot_courant:
            if premiere_ligne_lot is not None:
                # Total du lot précédent
                ws.merge_cells(f"A{ligne}:E{ligne}")
                ws.cell(row=ligne, column=1, value=f"Total {lot_courant}")
                _appliquer_style(ws.cell(row=ligne, column=1), styles["total"])
                cellule_total = ws.cell(row=ligne, column=6)
                cellule_total.value = f"=SUM(F{premiere_ligne_lot}:F{ligne - 1})"
                _appliquer_style(cellule_total, styles["total"])
                sommes_lots.append((premiere_ligne_lot, ligne))
                ligne += 1

            ws.merge_cells(f"A{ligne}:F{ligne}")
            cellule_lot = ws.cell(row=ligne, column=1, value=lot_libelle.upper())
            _appliquer_style(cellule_lot, styles["lot"])
            ws.row_dimensions[ligne].height = 18
            premiere_ligne_lot = ligne + 1
            lot_courant = lot_libelle
            chapitre_courant = None
            ligne += 1

        # Séparateur de chapitre
        chapitre = (article.chapitre or "").strip()
        if chapitre and chapitre != chapitre_courant:
            ws.merge_cells(f"A{ligne}:F{ligne}")
            cellule_chap = ws.cell(row=ligne, column=1, value=chapitre)
            _appliquer_style(cellule_chap, styles["chapitre"])
            ws.row_dimensions[ligne].height = 16
            chapitre_courant = chapitre
            ligne += 1

        # Ligne article
        num = ws.cell(row=ligne, column=1, value=article.numero_article)
        _appliquer_style(num, styles["numero"])

        desig = ws.cell(row=ligne, column=2, value=article.intitule)
        _appliquer_style(desig, styles["article"])

        unite = ws.cell(row=ligne, column=3, value="")
        _appliquer_style(unite, styles["article"])

        qte = ws.cell(row=ligne, column=4, value=None)
        _appliquer_style(qte, styles["montant"])

        pu = ws.cell(row=ligne, column=5, value=None)
        _appliquer_style(pu, styles["montant"])

        montant = ws.cell(row=ligne, column=6, value=f"=D{ligne}*E{ligne}")
        _appliquer_style(montant, styles["montant"])

        ws.row_dimensions[ligne].height = 16
        ligne += 1

    # Total dernier lot
    if premiere_ligne_lot is not None:
        ws.merge_cells(f"A{ligne}:E{ligne}")
        ws.cell(row=ligne, column=1, value=f"Total {lot_courant}")
        _appliquer_style(ws.cell(row=ligne, column=1), styles["total"])
        cellule_total = ws.cell(row=ligne, column=6)
        cellule_total.value = f"=SUM(F{premiere_ligne_lot}:F{ligne - 1})"
        _appliquer_style(cellule_total, styles["total"])
        sommes_lots.append((premiere_ligne_lot, ligne))
        ligne += 1

    # Total général
    if sommes_lots:
        ligne += 1
        ws.merge_cells(f"A{ligne}:E{ligne}")
        total_gen = ws.cell(row=ligne, column=1, value="TOTAL GÉNÉRAL HT")
        _appliquer_style(total_gen, {
            "font": Font(bold=True, size=10, color="FFFFFF"),
            "fill": PatternFill("solid", fgColor="1E3A5F"),
            "alignment": Alignment(horizontal="right", vertical="center"),
        })
        formule = "+".join(f"F{r_fin}" for _, r_fin in sommes_lots)
        cellule_tg = ws.cell(row=ligne, column=6, value=f"={formule}")
        _appliquer_style(cellule_tg, {
            "font": Font(bold=True, size=10, color="FFFFFF"),
            "fill": PatternFill("solid", fgColor="1E3A5F"),
            "alignment": Alignment(horizontal="right"),
            "number_format": "#,##0.00",
        })
        ws.row_dimensions[ligne].height = 22

    # Figer les volets (en-têtes)
    ws.freeze_panes = "A5"

    flux = BytesIO()
    wb.save(flux)
    return flux.getvalue()


def exporter_articles_bpu(piece: PieceEcrite) -> bytes:
    """
    Génère un fichier BPU (.xlsx) depuis les articles CCTP d'une pièce écrite.
    Chaque article CCTP = une ligne BPU avec :
      code, désignation complète (intitulé + corps), unité, prix unitaire.
    Le corps de l'article est inclus en note de bas de cellule pour le cahier des charges.
    """
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    articles = list(
        piece.articles.select_related("lot").order_by(
            "lot__ordre", "lot__code", "chapitre", "numero_article"
        )
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "BPU"

    styles = _styles_excel_dpgf(wb)

    projet = piece.projet
    ws.merge_cells("A1:E1")
    titre = ws["A1"]
    titre.value = f"BORDEREAU DES PRIX UNITAIRES — {piece.intitule.upper()}"
    _appliquer_style(titre, {
        "font": Font(bold=True, size=12, color="1E3A5F"),
        "alignment": Alignment(horizontal="center", vertical="center"),
    })
    ws.row_dimensions[1].height = 20

    ws.merge_cells("A2:E2")
    sous_titre = ws["A2"]
    sous_titre.value = (
        f"Projet : {projet.reference} — {projet.intitule}"
        f"  |  Date : {timezone.localtime().strftime('%d/%m/%Y')}"
    )
    _appliquer_style(sous_titre, {
        "font": Font(size=9, color="4A6FA5"),
        "alignment": Alignment(horizontal="center"),
    })
    ws.row_dimensions[2].height = 14

    entetes = ["N°", "DÉSIGNATION ET CAHIER DES CHARGES", "UNITÉ", "PRIX UNITAIRE HT (€)", "OBSERVATIONS"]
    largeurs = [12, 70, 10, 20, 25]
    for col, (libelle, largeur) in enumerate(zip(entetes, largeurs), start=1):
        cellule = ws.cell(row=4, column=col, value=libelle)
        _appliquer_style(cellule, styles["entete"])
        ws.column_dimensions[get_column_letter(col)].width = largeur
    ws.row_dimensions[4].height = 30

    ligne = 5
    lot_courant = None
    chapitre_courant = None

    for article in articles:
        lot_cctp = article.lot
        lot_libelle = f"{lot_cctp.code} — {lot_cctp.intitule}" if lot_cctp else "Sans lot"

        if lot_libelle != lot_courant:
            ws.merge_cells(f"A{ligne}:E{ligne}")
            cellule_lot = ws.cell(row=ligne, column=1, value=lot_libelle.upper())
            _appliquer_style(cellule_lot, styles["lot"])
            ws.row_dimensions[ligne].height = 18
            lot_courant = lot_libelle
            chapitre_courant = None
            ligne += 1

        chapitre = (article.chapitre or "").strip()
        if chapitre and chapitre != chapitre_courant:
            ws.merge_cells(f"A{ligne}:E{ligne}")
            cellule_chap = ws.cell(row=ligne, column=1, value=chapitre)
            _appliquer_style(cellule_chap, styles["chapitre"])
            ws.row_dimensions[ligne].height = 16
            chapitre_courant = chapitre
            ligne += 1

        # Pour le BPU : désignation = intitulé + corps de l'article (cahier des charges)
        corps_texte = _html_vers_texte(article.corps_article or "")
        designation_complete = article.intitule
        if corps_texte:
            designation_complete = f"{article.intitule}\n{corps_texte}"

        num = ws.cell(row=ligne, column=1, value=article.numero_article)
        _appliquer_style(num, styles["numero"])

        desig = ws.cell(row=ligne, column=2, value=designation_complete)
        _appliquer_style(desig, {**styles["article"], "alignment": Alignment(
            horizontal="left", vertical="top", wrap_text=True
        )})

        unite = ws.cell(row=ligne, column=3, value="")
        _appliquer_style(unite, styles["article"])

        pu = ws.cell(row=ligne, column=4, value=None)
        _appliquer_style(pu, styles["montant"])

        obs = ws.cell(row=ligne, column=5, value="")
        _appliquer_style(obs, styles["article"])

        # Hauteur proportionnelle au nombre de lignes du corps
        nb_lignes = max(1, len(designation_complete.splitlines()))
        ws.row_dimensions[ligne].height = max(20, min(120, nb_lignes * 13))
        ligne += 1

    ws.freeze_panes = "A5"

    flux = BytesIO()
    wb.save(flux)
    return flux.getvalue()
