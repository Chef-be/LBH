"""Intégration bureautique WOPI / Collabora pour les modèles de documents."""

from __future__ import annotations

from io import BytesIO
from urllib import parse as urllib_parse
from urllib import request as urllib_request
from xml.etree import ElementTree

from django.conf import settings
from django.core.cache import cache
from django.core.files.base import ContentFile
from django.core.signing import BadSignature, SignatureExpired, TimestampSigner
from django.utils import timezone
from django.utils.text import slugify
from docx import Document as DocumentWord
from openpyxl import Workbook

from .models import ModeleDocument

SALT_WOPI_MODELES = "pieces-ecrites-wopi-modeles"
WOPI_LOCK_TIMEOUT = 30 * 60
WOPI_TOKEN_TIMEOUT = 8 * 60 * 60
TYPES_TABLEUR = {"bpu", "dpgf", "dqe"}


def type_bureautique_modele(modele: ModeleDocument) -> str:
    return "tableur" if modele.type_document in TYPES_TABLEUR else "texte"


def extension_gabarit_modele(modele: ModeleDocument) -> str:
    return ".xlsx" if type_bureautique_modele(modele) == "tableur" else ".docx"


def type_mime_gabarit_modele(modele: ModeleDocument) -> str:
    if extension_gabarit_modele(modele) == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _nom_fichier_gabarit(modele: ModeleDocument) -> str:
    base = slugify(modele.code or modele.libelle or f"modele-{modele.pk}") or "modele"
    return f"{base}{extension_gabarit_modele(modele)}"


def _contenu_gabarit_docx(modele: ModeleDocument) -> bytes:
    document = DocumentWord()
    document.add_heading(modele.libelle or "Modèle de document", level=1)
    document.add_paragraph("Projet : {reference_projet} — {nom_projet}")
    document.add_paragraph("Maître d'ouvrage : {maitre_ouvrage}")
    document.add_paragraph("Rédacteur : {redacteur_nom}")
    document.add_paragraph("Date : {date_generation}")
    document.add_paragraph(" ")
    document.add_paragraph("Rédigez votre modèle ici. Vous pouvez insérer des variables comme {lot_intitule}.")
    flux = BytesIO()
    document.save(flux)
    return flux.getvalue()


def _contenu_gabarit_xlsx(modele: ModeleDocument) -> bytes:
    classeur = Workbook()
    feuille = classeur.active
    feuille.title = "Modele"
    feuille["A1"] = "Projet"
    feuille["B1"] = "{reference_projet} - {nom_projet}"
    feuille["A2"] = "Lot"
    feuille["B2"] = "{lot_numero} - {lot_intitule}"
    feuille["A4"] = "Code"
    feuille["B4"] = "Désignation"
    feuille["C4"] = "Unité"
    feuille["D4"] = "Quantité"
    feuille["E4"] = "Prix unitaire"
    feuille["F4"] = "Montant"
    feuille["A5"] = "{reference_projet}"
    feuille["B5"] = "Ligne de démonstration"
    feuille["C5"] = "u"
    feuille["D5"] = 1
    feuille["E5"] = 0
    feuille["F5"] = "=D5*E5"
    feuille.freeze_panes = "A5"
    flux = BytesIO()
    classeur.save(flux)
    return flux.getvalue()


def assurer_gabarit_bureautique(modele: ModeleDocument) -> ModeleDocument:
    extension_attendue = extension_gabarit_modele(modele)
    if modele.gabarit and modele.gabarit.name.lower().endswith(extension_attendue):
        return modele

    contenu = _contenu_gabarit_xlsx(modele) if extension_attendue == ".xlsx" else _contenu_gabarit_docx(modele)
    modele.gabarit.save(_nom_fichier_gabarit(modele), ContentFile(contenu), save=True)
    return modele


def lire_contenu_gabarit(modele: ModeleDocument) -> bytes:
    assurer_gabarit_bureautique(modele)
    with modele.gabarit.open("rb") as fichier:
        return fichier.read()


def enregistrer_contenu_gabarit(modele: ModeleDocument, contenu: bytes) -> ModeleDocument:
    assurer_gabarit_bureautique(modele)
    modele.gabarit.save(_nom_fichier_gabarit(modele), ContentFile(contenu), save=True)
    return modele


def _signer() -> TimestampSigner:
    return TimestampSigner(salt=SALT_WOPI_MODELES)


def creer_jeton_wopi_modele(modele: ModeleDocument, utilisateur) -> str:
    payload = f"{modele.pk}:{utilisateur.pk}:{int(timezone.now().timestamp())}"
    return _signer().sign(payload)


def verifier_jeton_wopi_modele(modele: ModeleDocument, jeton: str) -> dict[str, str]:
    if not jeton:
        raise PermissionError("Jeton WOPI manquant.")

    try:
        valeur = _signer().unsign(jeton, max_age=WOPI_TOKEN_TIMEOUT)
    except SignatureExpired as exc:
        raise PermissionError("Jeton WOPI expiré.") from exc
    except BadSignature as exc:
        raise PermissionError("Jeton WOPI invalide.") from exc

    modele_id, utilisateur_id, _horodatage = (valeur.split(":", 2) + ["", ""])[:3]
    if modele_id != str(modele.pk):
        raise PermissionError("Jeton WOPI invalide pour ce modèle.")
    return {"modele_id": modele_id, "utilisateur_id": utilisateur_id}


def cle_verrou_modele(modele: ModeleDocument) -> str:
    return f"pieces-ecrites:wopi-lock:{modele.pk}"


def lire_verrou_modele(modele: ModeleDocument) -> str | None:
    return cache.get(cle_verrou_modele(modele))


def definir_verrou_modele(modele: ModeleDocument, valeur: str):
    cache.set(cle_verrou_modele(modele), valeur, WOPI_LOCK_TIMEOUT)


def supprimer_verrou_modele(modele: ModeleDocument):
    cache.delete(cle_verrou_modele(modele))


def _discovery_urlsrc(extension: str) -> str:
    """
    Interroge la discovery WOPI via l'URL interne Docker, puis remplace le préfixe
    interne (https://lbh-collabora:9980) par l'URL publique accessible depuis le navigateur.
    """
    url_interne = getattr(settings, "COLLABORA_URL", "http://lbh-collabora:9980").rstrip("/")
    url_publique = getattr(settings, "COLLABORA_PUBLIC_URL", "https://lbh-economiste.com/collabora").rstrip("/")

    with urllib_request.urlopen(f"{url_interne}/hosting/discovery", timeout=10) as reponse:
        contenu = reponse.read()

    racine = ElementTree.fromstring(contenu)
    candidats = []
    for action in racine.findall(".//action"):
        ext = action.attrib.get("ext")
        nom = action.attrib.get("name")
        urlsrc_brut = action.attrib.get("urlsrc", "")
        if not (ext == extension and urlsrc_brut):
            continue
        # Remplacer le préfixe interne par l'URL publique Nginx
        # ex : https://lbh-collabora:9980/browser/... → https://lbh-economiste.com/collabora/browser/...
        urlsrc = urlsrc_brut.replace("https://lbh-collabora:9980", url_publique)
        urlsrc = urlsrc.replace("http://lbh-collabora:9980", url_publique)
        if nom == "edit":
            return urlsrc
        candidats.append(urlsrc)
    if candidats:
        return candidats[0]
    raise RuntimeError(f"Aucune action Collabora disponible pour l'extension {extension}.")


def construire_url_editeur_collabora(wopi_src: str, access_token: str, extension: str) -> str:
    urlsrc = _discovery_urlsrc(extension.lstrip("."))
    parametres = (
        f"WOPISrc={urllib_parse.quote(wopi_src, safe='')}"
        "&lang=fr"
    )
    return f"{urlsrc}{parametres}"


def nom_affichage_gabarit(modele: ModeleDocument) -> str:
    if modele.gabarit and getattr(modele.gabarit, "name", ""):
        return modele.gabarit.name.rsplit("/", 1)[-1]

    nom = modele.libelle or _nom_fichier_gabarit(modele)
    extension = extension_gabarit_modele(modele)
    if not nom.lower().endswith(extension):
        return f"{slugify(nom) or 'modele'}{extension}"
    return nom
