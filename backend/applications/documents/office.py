"""Intégration WOPI / Collabora pour les documents GED projet."""

from __future__ import annotations

from io import BytesIO
from urllib import parse as urllib_parse
from urllib import request as urllib_request
from xml.etree import ElementTree

from django.conf import settings
from django.core.cache import cache
from django.core.files.base import ContentFile
from django.core.signing import BadSignature, SignatureExpired, TimestampSigner
from django.utils import timezone

from docx import Document as DocumentWord
from openpyxl import Workbook

from .models import Document
from .services import est_document_bureautique_editable

SALT_WOPI_DOCUMENTS = "documents-wopi-ged"
WOPI_LOCK_TIMEOUT = 30 * 60
WOPI_TOKEN_TIMEOUT = 8 * 60 * 60


def extension_bureautique_document(document: Document) -> str:
    nom = document.nom_fichier_origine or ""
    extension = f".{nom.rsplit('.', 1)[-1].lower()}" if "." in nom else ""
    if extension in {".doc", ".docx", ".odt", ".xls", ".xlsx", ".xlsm", ".ods"}:
        return extension
    if (document.type_mime or "") in {
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.oasis.opendocument.spreadsheet",
    }:
        return ".xlsx"
    return ".docx"


def type_bureautique_document(document: Document) -> str:
    return "tableur" if extension_bureautique_document(document) in {".xls", ".xlsx", ".xlsm", ".ods"} else "texte"


def type_mime_bureautique_document(document: Document) -> str:
    extension = extension_bureautique_document(document)
    if extension in {".xls", ".xlsx", ".xlsm", ".ods"}:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def nom_affichage_document(document: Document) -> str:
    nom = (document.nom_fichier_origine or "").strip()
    if nom:
        return nom
    base = (document.intitule or document.reference or "document").strip()
    return f"{base}{extension_bureautique_document(document)}"


def _contenu_document_docx(document: Document) -> bytes:
    fichier = DocumentWord()
    fichier.add_heading(document.intitule or document.reference or "Document projet", level=1)
    if document.projet_id:
        fichier.add_paragraph(f"Projet : {document.projet.reference} — {document.projet.intitule}")
    fichier.add_paragraph("Rédigez votre document ici.")
    flux = BytesIO()
    fichier.save(flux)
    return flux.getvalue()


def _contenu_document_xlsx(document: Document) -> bytes:
    classeur = Workbook()
    feuille = classeur.active
    feuille.title = "Document"
    feuille["A1"] = "Projet"
    feuille["B1"] = f"{document.projet.reference} - {document.projet.intitule}" if document.projet_id else ""
    feuille["A3"] = "Élément"
    feuille["B3"] = "Valeur"
    feuille["A4"] = "Référence"
    feuille["B4"] = document.reference
    feuille["A5"] = "Intitulé"
    feuille["B5"] = document.intitule
    flux = BytesIO()
    classeur.save(flux)
    return flux.getvalue()


def assurer_fichier_bureautique_document(document: Document) -> Document:
    if document.fichier and est_document_bureautique_editable(document.nom_fichier_origine, document.type_mime):
        return document

    extension = extension_bureautique_document(document)
    contenu = _contenu_document_xlsx(document) if extension in {".xls", ".xlsx", ".xlsm", ".ods"} else _contenu_document_docx(document)
    nom = nom_affichage_document(document)
    document.fichier.save(nom, ContentFile(contenu), save=False)
    document.nom_fichier_origine = nom
    document.type_mime = type_mime_bureautique_document(document)
    document.taille_octets = len(contenu)
    document.save(update_fields=["fichier", "nom_fichier_origine", "type_mime", "taille_octets", "date_modification"])
    return document


def lire_contenu_document(document: Document) -> bytes:
    assurer_fichier_bureautique_document(document)
    with document.fichier.open("rb") as fichier:
        return fichier.read()


def enregistrer_contenu_document(document: Document, contenu: bytes) -> Document:
    assurer_fichier_bureautique_document(document)
    document.fichier.save(nom_affichage_document(document), ContentFile(contenu), save=False)
    document.taille_octets = len(contenu)
    document.save(update_fields=["fichier", "taille_octets", "date_modification"])
    return document


def _signer() -> TimestampSigner:
    return TimestampSigner(salt=SALT_WOPI_DOCUMENTS)


def creer_jeton_wopi_document(document: Document, utilisateur) -> str:
    payload = f"{document.pk}:{utilisateur.pk}:{int(timezone.now().timestamp())}"
    return _signer().sign(payload)


def verifier_jeton_wopi_document(document: Document, jeton: str) -> dict[str, str]:
    if not jeton:
        raise PermissionError("Jeton WOPI manquant.")
    try:
        valeur = _signer().unsign(jeton, max_age=WOPI_TOKEN_TIMEOUT)
    except SignatureExpired as exc:
        raise PermissionError("Jeton WOPI expiré.") from exc
    except BadSignature as exc:
        raise PermissionError("Jeton WOPI invalide.") from exc
    document_id, utilisateur_id, _horodatage = (valeur.split(":", 2) + ["", ""])[:3]
    if document_id != str(document.pk):
        raise PermissionError("Jeton WOPI invalide pour ce document.")
    return {"document_id": document_id, "utilisateur_id": utilisateur_id}


def cle_verrou_document(document: Document) -> str:
    return f"documents:wopi-lock:{document.pk}"


def lire_verrou_document(document: Document) -> str | None:
    return cache.get(cle_verrou_document(document))


def definir_verrou_document(document: Document, valeur: str):
    cache.set(cle_verrou_document(document), valeur, WOPI_LOCK_TIMEOUT)


def supprimer_verrou_document(document: Document):
    cache.delete(cle_verrou_document(document))


def _discovery_urlsrc(extension: str) -> str:
    """
    Interroge la discovery WOPI via l'URL interne Docker, puis remplace le préfixe
    interne par l'URL publique accessible depuis le navigateur.
    """
    url_interne = getattr(settings, "COLLABORA_URL", "http://lbh-collabora:9980").rstrip("/")
    url_publique = getattr(settings, "COLLABORA_PUBLIC_URL", "https://lbh-economiste.com/collabora").rstrip("/")

    with urllib_request.urlopen(f"{url_interne}/hosting/discovery", timeout=10) as reponse:
        contenu = reponse.read()
    racine = ElementTree.fromstring(contenu)
    candidats = []
    for action in racine.findall(".//action"):
        ext = action.attrib.get("ext")
        nom = action.attrib.get("name")
        urlsrc_brut = action.attrib.get("urlsrc", "")
        if not (ext == extension and urlsrc_brut):
            continue
        urlsrc = urlsrc_brut.replace("https://lbh-collabora:9980", url_publique)
        urlsrc = urlsrc.replace("http://lbh-collabora:9980", url_publique)
        if nom == "edit":
            return urlsrc
        candidats.append(urlsrc)
    if candidats:
        return candidats[0]
    raise RuntimeError(f"Aucune action Collabora disponible pour l'extension {extension}.")


def construire_url_editeur_collabora_document(wopi_src: str, access_token: str, extension: str) -> str:
    # access_token est inclus dans l'URL afin que le JS de cool.html puisse le lire et
    # l'insérer dans le chemin WebSocket (/cool/{WOPISrc+token}/ws).
    # Il est AUSSI transmis dans le corps du formulaire HTML (POST) pour que Collabora
    # l'utilise lors des appels GetFile/PutFile.
    urlsrc = _discovery_urlsrc(extension.lstrip("."))
    token_encode = urllib_parse.quote(access_token, safe="") if access_token else ""
    return f"{urlsrc}WOPISrc={urllib_parse.quote(wopi_src, safe='')}&lang=fr&access_token={token_encode}&access_token_ttl=28800000"
