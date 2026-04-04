"""Services métier pour l'économie, les profils de main-d'œuvre et les simulateurs."""

from __future__ import annotations

from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
import subprocess
import tempfile
from statistics import mean

from docx import Document as DocumentWord
from docx.shared import Pt

SMIC_HORAIRE_2026 = {
    "metropole": Decimal("12.02"),
    "dom": Decimal("12.02"),
    "mayotte": Decimal("9.33"),
}

LIBELLES_LOCALISATIONS = {
    "nationale": "Nationale",
    "metropole": "Métropole",
    "dom": "Autre DOM",
    "mayotte": "Mayotte",
}


def d(valeur) -> Decimal:
    if valeur is None or valeur == "":
        return Decimal("0")
    return Decimal(str(valeur))


def arrondir(valeur: Decimal, decimales: int = 2) -> Decimal:
    quantification = Decimal(f"0.{'0' * decimales}")
    return valeur.quantize(quantification, rounding=ROUND_HALF_UP)


def calculer_indicateurs_commerciaux_etude_prix(etude) -> dict[str, Decimal]:
    debourse = d(getattr(etude, "debourse_sec_ht", 0))
    taux_frais_chantier = d(getattr(etude, "taux_frais_chantier", 0))
    taux_frais_generaux = d(getattr(etude, "taux_frais_generaux", 0))
    taux_aleas = d(getattr(etude, "taux_aleas", 0))
    taux_marge_cible = d(getattr(etude, "taux_marge_cible", 0))

    montant_frais_chantier = debourse * taux_frais_chantier
    montant_frais_generaux = (debourse + montant_frais_chantier) * taux_frais_generaux
    montant_aleas = debourse * taux_aleas
    cout_revient = debourse + montant_frais_chantier + montant_frais_generaux + montant_aleas
    marge_previsionnelle = cout_revient * taux_marge_cible
    prix_vente = cout_revient + marge_previsionnelle
    coefficient_k = prix_vente / debourse if debourse else Decimal("0")

    return {
        "montant_frais_chantier_ht": arrondir(montant_frais_chantier, 4),
        "montant_frais_generaux_ht": arrondir(montant_frais_generaux, 4),
        "montant_aleas_ht": arrondir(montant_aleas, 4),
        "cout_revient_ht": arrondir(cout_revient, 4),
        "marge_previsionnelle_ht": arrondir(marge_previsionnelle, 4),
        "prix_vente_ht": arrondir(prix_vente, 4),
        "coefficient_k": arrondir(coefficient_k, 4),
        "seuil_rentabilite_ht": arrondir(cout_revient, 4),
    }


def proposer_achats_depuis_etude_prix(etude, *, remplacer: bool = False) -> list:
    from .models import AchatEtudePrix

    if remplacer:
        etude.achats.all().delete()

    achats = []
    ordre_depart = etude.achats.count()
    lignes = list(etude.lignes.filter(type_ressource="matiere").order_by("ordre", "designation"))
    for index, ligne in enumerate(lignes, start=1):
        achat, _ = AchatEtudePrix.objects.update_or_create(
            etude=etude,
            ligne_source=ligne,
            defaults={
                "ordre": ordre_depart + index,
                "designation": ligne.designation,
                "unite_achat": ligne.unite or "u",
                "quantite_besoin": ligne.quantite,
                "quantite_conditionnement": ligne.quantite if d(ligne.quantite) > 0 else Decimal("1"),
                "prix_unitaire_achat_ht": ligne.cout_unitaire_ht,
                "observations": ligne.observations,
            },
        )
        achats.append(achat)
    return achats


def construire_cadrage_etude_prix(etude) -> dict[str, object]:
    from applications.projets.services import construire_processus_recommande

    projet = getattr(etude, "projet", None)
    processus = construire_processus_recommande(projet=projet) if projet else {
        "clientele": {"code": "autre", "libelle": "Autre contexte"},
        "objectif": {"code": "autre", "libelle": "Autre objectif"},
        "questions_ouverture": _questions_simulateur(),
        "methodes_estimation": [],
        "livrables_prioritaires": [],
        "indicateurs_clefs": [],
        "points_de_controle": [],
    }

    lignes = list(etude.lignes.order_by("-montant_ht", "ordre"))
    postes_sensibles = [
        {
            "id": str(ligne.id),
            "designation": ligne.designation,
            "type_ressource": ligne.type_ressource,
            "montant_ht": arrondir(d(ligne.montant_ht), 4),
        }
        for ligne in lignes[:5]
    ]
    total_achats = etude.achats.count() if hasattr(etude, "achats") else 0
    couverture_analytique = Decimal("0")
    if d(etude.debourse_sec_ht) > 0 and lignes:
        couverture_analytique = Decimal("1")

    return {
        "clientele": processus.get("clientele"),
        "objectif": processus.get("objectif"),
        "resume": "Cadrage métier issu du type de clientèle, de l'objectif de mission et des méthodes d'estimation recommandées.",
        "questions": processus.get("questions_ouverture", []),
        "methodes_recommandees": processus.get("methodes_estimation", []),
        "livrables_recommandes": processus.get("livrables_prioritaires", []),
        "indicateurs_prioritaires": processus.get("indicateurs_clefs", []),
        "points_vigilance": processus.get("points_de_controle", []),
        "postes_sensibles": postes_sensibles,
        "couverture_analytique": arrondir(couverture_analytique, 4),
        "achats_prepares": total_achats,
        "sous_detail_disponible": bool(lignes),
    }


def _surface_reference_projet(projet) -> Decimal:
    from applications.batiment.models import ProgrammeBatiment

    surfaces = []
    for programme in ProgrammeBatiment.objects.filter(projet=projet):
        for valeur in (
            getattr(programme, "shon_totale", None),
            getattr(programme, "shab_totale", None),
            getattr(programme, "emprise_sol", None),
        ):
            if valeur:
                surfaces.append(d(valeur))
                break
    return sum(surfaces, Decimal("0"))


def _indicateurs_metres_projet(projet) -> dict[str, Decimal | int]:
    from applications.metres.models import LigneMetre, Metre

    metres = Metre.objects.filter(projet=projet).prefetch_related("lignes")
    montant_total = Decimal("0")
    nombre_metres = metres.count()
    nombre_lignes = 0
    for metre in metres:
        montant_total += d(getattr(metre, "montant_total_ht", 0))
        nombre_lignes += metre.lignes.count()

    total_quantites = Decimal("0")
    for ligne in LigneMetre.objects.filter(metre__projet=projet):
        total_quantites += d(ligne.quantite)

    return {
        "nombre_metres": nombre_metres,
        "nombre_lignes_metre": nombre_lignes,
        "montant_metres_ht": montant_total,
        "total_quantites_metrees": total_quantites,
    }


def _indicateurs_programme_batiment(projet) -> dict[str, object]:
    from applications.batiment.models import LocalProgramme, ProgrammeBatiment

    programmes = list(ProgrammeBatiment.objects.filter(projet=projet))
    types_operation = []
    types_batiment = []
    categories_locaux = []
    for programme in programmes:
        if programme.type_operation:
            types_operation.append(programme.type_operation)
        if programme.type_batiment:
            types_batiment.append(programme.type_batiment)

    for local in LocalProgramme.objects.filter(programme__projet=projet):
        if local.categorie:
            categories_locaux.append(local.categorie)

    return {
        "nombre_programmes_batiment": len(programmes),
        "types_operation_batiment": sorted(set(types_operation)),
        "types_batiment": sorted(set(types_batiment)),
        "categories_locaux": sorted(set(categories_locaux)),
    }


def _montant_reference_projet(projet) -> Decimal:
    montant = d(getattr(projet, "montant_marche", None))
    if montant > 0:
        return montant
    montant = d(getattr(projet, "montant_estime", None))
    if montant > 0:
        return montant
    lots = getattr(projet, "lots", None)
    if lots is not None:
        total_lots = sum((d(lot.montant_estime) for lot in lots.all() if lot.montant_estime), Decimal("0"))
        if total_lots > 0:
            return total_lots
    metres = getattr(projet, "metres", None)
    if metres is not None:
        total_metres = sum((d(metre.montant_total_ht) for metre in metres.all()), Decimal("0"))
        if total_metres > 0:
            return total_metres
    return Decimal("0")


def _profil_reference_projet(projet) -> dict[str, object]:
    indicateurs_metres = _indicateurs_metres_projet(projet)
    indicateurs_batiment = _indicateurs_programme_batiment(projet)
    return {
        "id": str(projet.id),
        "reference": projet.reference,
        "intitule": projet.intitule,
        "type_projet": projet.type_projet,
        "clientele_cible": getattr(projet, "clientele_cible", ""),
        "objectif_mission": getattr(projet, "objectif_mission", ""),
        "phase_actuelle": projet.phase_actuelle or "",
        "commune": projet.commune or "",
        "departement": projet.departement or "",
        "surface_reference": _surface_reference_projet(projet),
        "montant_reference": _montant_reference_projet(projet),
        "nombre_lots": getattr(projet, "lots", None).count() if hasattr(projet, "lots") else 0,
        **indicateurs_metres,
        **indicateurs_batiment,
    }


def _score_similarite_profils(reference: dict[str, object], candidat: dict[str, object]) -> Decimal:
    score = Decimal("0")

    if reference["type_projet"] == candidat["type_projet"]:
        score += Decimal("0.28")
    if reference["clientele_cible"] == candidat["clientele_cible"]:
        score += Decimal("0.20")
    if reference["objectif_mission"] == candidat["objectif_mission"]:
        score += Decimal("0.12")
    if reference["phase_actuelle"] == candidat["phase_actuelle"]:
        score += Decimal("0.08")
    if reference["departement"] and reference["departement"] == candidat["departement"]:
        score += Decimal("0.08")
    if reference["commune"] and reference["commune"] == candidat["commune"]:
        score += Decimal("0.04")

    types_batiment_ref = set(reference.get("types_batiment", []) or [])
    types_batiment_cand = set(candidat.get("types_batiment", []) or [])
    if types_batiment_ref and types_batiment_cand:
        if types_batiment_ref & types_batiment_cand:
            score += Decimal("0.05")

    types_operation_ref = set(reference.get("types_operation_batiment", []) or [])
    types_operation_cand = set(candidat.get("types_operation_batiment", []) or [])
    if types_operation_ref and types_operation_cand:
        if types_operation_ref & types_operation_cand:
            score += Decimal("0.04")

    categories_ref = set(reference.get("categories_locaux", []) or [])
    categories_cand = set(candidat.get("categories_locaux", []) or [])
    if categories_ref and categories_cand:
        ratio_categories = Decimal(len(categories_ref & categories_cand)) / Decimal(max(len(categories_ref), len(categories_cand)))
        score += Decimal("0.03") * ratio_categories

    surface_ref = d(reference["surface_reference"])
    surface_cand = d(candidat["surface_reference"])
    if surface_ref > 0 and surface_cand > 0:
        ecart_surface = abs(surface_ref - surface_cand) / max(surface_ref, surface_cand)
        score += max(Decimal("0"), Decimal("0.12") * (Decimal("1") - min(ecart_surface, Decimal("1"))))

    montant_ref = d(reference["montant_reference"])
    montant_cand = d(candidat["montant_reference"])
    if montant_ref > 0 and montant_cand > 0:
        ecart_montant = abs(montant_ref - montant_cand) / max(montant_ref, montant_cand)
        score += max(Decimal("0"), Decimal("0.08") * (Decimal("1") - min(ecart_montant, Decimal("1"))))

    montant_metre_ref = d(reference.get("montant_metres_ht", 0))
    montant_metre_cand = d(candidat.get("montant_metres_ht", 0))
    if montant_metre_ref > 0 and montant_metre_cand > 0:
        ecart_metres = abs(montant_metre_ref - montant_metre_cand) / max(montant_metre_ref, montant_metre_cand)
        score += max(Decimal("0"), Decimal("0.04") * (Decimal("1") - min(ecart_metres, Decimal("1"))))

    lignes_metre_ref = Decimal(str(reference.get("nombre_lignes_metre", 0) or 0))
    lignes_metre_cand = Decimal(str(candidat.get("nombre_lignes_metre", 0) or 0))
    if lignes_metre_ref > 0 and lignes_metre_cand > 0:
        ecart_lignes = abs(lignes_metre_ref - lignes_metre_cand) / max(lignes_metre_ref, lignes_metre_cand)
        score += max(Decimal("0"), Decimal("0.02") * (Decimal("1") - min(ecart_lignes, Decimal("1"))))

    return arrondir(min(score, Decimal("1")), 4)


def calculer_comparatif_estimation_etude_prix(etude) -> dict[str, object]:
    from applications.projets.models import Projet

    projet = getattr(etude, "projet", None)
    profil_projet = _profil_reference_projet(projet) if projet else {
        "id": "",
        "reference": "",
        "intitule": etude.intitule,
        "type_projet": "",
        "clientele_cible": "",
        "objectif_mission": "",
        "phase_actuelle": "",
        "commune": "",
        "departement": "",
        "surface_reference": Decimal("0"),
        "montant_reference": Decimal("0"),
        "nombre_lots": 0,
        "nombre_metres": 0,
        "nombre_lignes_metre": 0,
        "montant_metres_ht": Decimal("0"),
        "total_quantites_metrees": Decimal("0"),
        "nombre_programmes_batiment": 0,
        "types_operation_batiment": [],
        "types_batiment": [],
        "categories_locaux": [],
    }

    candidats = []
    queryset = Projet.objects.exclude(id=getattr(projet, "id", None)).prefetch_related("lots", "metres")
    for candidat in queryset:
        profil_candidat = _profil_reference_projet(candidat)
        if d(profil_candidat["montant_reference"]) <= 0:
            continue
        score = _score_similarite_profils(profil_projet, profil_candidat)
        if score <= 0:
            continue
        candidats.append((score, profil_candidat))

    candidats.sort(key=lambda item: (item[0], item[1]["montant_reference"]), reverse=True)
    top = candidats[:5]

    taux_similarite_moyen = arrondir(
        Decimal(str(mean([float(score) for score, _ in top]))) if top else Decimal("0"),
        4,
    )

    estimation_ratio = Decimal("0")
    surface_courante = d(profil_projet["surface_reference"])
    projets_similaires = []
    if top:
        if surface_courante > 0:
            ratios = []
            poids = Decimal("0")
            for score, profil in top:
                surface = d(profil["surface_reference"])
                montant = d(profil["montant_reference"])
                if surface > 0 and montant > 0:
                    ratios.append((score, montant / surface))
                    poids += score
            if ratios and poids > 0:
                estimation_ratio = sum((score * ratio for score, ratio in ratios), Decimal("0")) / poids
                estimation_ratio *= surface_courante
        if estimation_ratio <= 0:
            poids = sum((score for score, _ in top), Decimal("0"))
            if poids > 0:
                estimation_ratio = sum((score * d(profil["montant_reference"]) for score, profil in top), Decimal("0")) / poids

        for score, profil in top:
            projets_similaires.append(
                {
                    **profil,
                    "score_similarite": arrondir(score, 4),
                }
            )

    prix_historiques_lot = []
    try:
        from .models import EtudePrix

        etudes_historiques = EtudePrix.objects.filter(
            lot_type=etude.lot_type,
            statut__in=["validee", "publiee"],
        ).exclude(id=etude.id)
        prix_historiques_lot = [d(item.prix_vente_ht) for item in etudes_historiques if d(item.prix_vente_ht) > 0]
    except Exception:
        prix_historiques_lot = []

    estimation_rex = (
        arrondir(Decimal(str(mean([float(valeur) for valeur in prix_historiques_lot]))), 4)
        if prix_historiques_lot else Decimal("0")
    )
    estimation_analytique = d(getattr(etude, "prix_vente_ht", 0))

    alertes = []
    if estimation_ratio > 0 and estimation_analytique > 0:
        ecart = abs(estimation_analytique - estimation_ratio) / max(estimation_analytique, estimation_ratio)
        if ecart > Decimal("0.15"):
            alertes.append("Écart significatif entre l'estimation ratio et l'estimation analytique.")
    if estimation_rex > 0 and estimation_analytique > 0:
        ecart = abs(estimation_analytique - estimation_rex) / max(estimation_analytique, estimation_rex)
        if ecart > Decimal("0.15"):
            alertes.append("Écart significatif entre le retour d'expérience et l'estimation analytique.")
    if not projets_similaires:
        alertes.append("Aucun projet comparable suffisamment renseigné n'a été trouvé dans la base.")

    return {
        "profil_projet": profil_projet,
        "projets_similaires": projets_similaires,
        "taux_similarite_moyen": taux_similarite_moyen,
        "estimation_ratio_ht": arrondir(estimation_ratio, 4),
        "estimation_rex_ht": arrondir(estimation_rex, 4),
        "estimation_analytique_ht": arrondir(estimation_analytique, 4),
        "ecart_ratio_vs_analytique_ht": arrondir(estimation_analytique - estimation_ratio, 4),
        "ecart_rex_vs_analytique_ht": arrondir(estimation_analytique - estimation_rex, 4),
        "alertes": alertes,
    }


def _ajouter_tableau_docx(document, titre: str, lignes: list[tuple[str, str]]) -> None:
    document.add_heading(titre, level=2)
    tableau = document.add_table(rows=0, cols=2)
    tableau.style = "Table Grid"
    for libelle, valeur in lignes:
        cellules = tableau.add_row().cells
        cellules[0].text = libelle
        cellules[1].text = valeur


def generer_docx_livrable_estimation(etude, comparatif: dict[str, object], *, mode: str) -> bytes:
    document = DocumentWord()
    titre = document.add_paragraph()
    run_titre = titre.add_run(
        "Note de vérification d'enveloppe"
        if mode == "moa"
        else "Note d'estimation consolidée"
    )
    run_titre.bold = True
    run_titre.font.size = Pt(16)

    sous_titre = document.add_paragraph(
        f"{etude.code or 'Étude de prix'} — {etude.intitule}"
    )
    sous_titre.runs[0].font.size = Pt(9)

    document.add_paragraph(
        "Ce livrable consolide l'estimation analytique, la vérification par ratio et les retours d'expérience internes."
    )

    if mode == "moa":
        document.add_heading("Positionnement maîtrise d'ouvrage", level=1)
        document.add_paragraph(
            "La note vise à apprécier l'adéquation de l'enveloppe financière du projet avec son programme, "
            "son niveau de définition et les références comparables disponibles."
        )
    else:
        document.add_heading("Positionnement maîtrise d'œuvre", level=1)
        document.add_paragraph(
            "La note vise à consolider l'estimation par ouvrages élémentaires, à comparer les méthodes "
            "de chiffrage et à sécuriser la cohérence des pièces DCE."
        )

    profil = comparatif["profil_projet"]
    _ajouter_tableau_docx(
        document,
        "Données de base du projet",
        [
            ("Projet", f"{profil.get('reference') or ''} — {profil.get('intitule') or ''}".strip(" —")),
            ("Surface de référence", f"{float(profil.get('surface_reference') or 0):.2f} m²"),
            ("Montant de référence", f"{float(profil.get('montant_reference') or 0):,.2f} €".replace(",", " ")),
            ("Montant des métrés", f"{float(profil.get('montant_metres_ht') or 0):,.2f} €".replace(",", " ")),
            ("Lignes de métré", str(profil.get("nombre_lignes_metre") or 0)),
            ("Types bâtiment", ", ".join(profil.get("types_batiment") or []) or "—"),
        ],
    )

    _ajouter_tableau_docx(
        document,
        "Comparaison des estimations",
        [
            ("Estimation par ratio", f"{float(comparatif['estimation_ratio_ht']):,.2f} €".replace(",", " ")),
            ("Retour d'expérience", f"{float(comparatif['estimation_rex_ht']):,.2f} €".replace(",", " ")),
            ("Estimation analytique", f"{float(comparatif['estimation_analytique_ht']):,.2f} €".replace(",", " ")),
            ("Seuil de rentabilité", f"{float(etude.seuil_rentabilite_ht):,.2f} €".replace(",", " ")),
            ("Prix de vente prévisionnel", f"{float(etude.prix_vente_ht):,.2f} €".replace(",", " ")),
            ("Taux moyen de similarité", f"{float(comparatif['taux_similarite_moyen']) * 100:.2f} %"),
        ],
    )

    document.add_heading("Projets similaires mobilisés", level=2)
    if comparatif["projets_similaires"]:
        tableau = document.add_table(rows=1, cols=5)
        tableau.style = "Table Grid"
        entete = tableau.rows[0].cells
        entete[0].text = "Référence"
        entete[1].text = "Projet"
        entete[2].text = "Montant"
        entete[3].text = "Surface"
        entete[4].text = "Similarité"
        for projet in comparatif["projets_similaires"]:
            ligne = tableau.add_row().cells
            ligne[0].text = projet.get("reference") or ""
            ligne[1].text = projet.get("intitule") or ""
            ligne[2].text = f"{float(projet.get('montant_reference') or 0):,.2f} €".replace(",", " ")
            ligne[3].text = f"{float(projet.get('surface_reference') or 0):.2f} m²"
            ligne[4].text = f"{float(projet.get('score_similarite') or 0) * 100:.2f} %"
    else:
        document.add_paragraph("Aucun projet comparable suffisamment renseigné n'a été trouvé.")

    document.add_heading("Conclusion et points de vigilance", level=2)
    if comparatif["alertes"]:
        for alerte in comparatif["alertes"]:
            document.add_paragraph(alerte, style="List Bullet")
    else:
        document.add_paragraph("Aucune alerte majeure n'a été détectée sur la convergence des méthodes d'estimation.")

    if mode == "moa":
        document.add_paragraph(
            "Recommandation MOA : consolider l'enveloppe de travaux à partir de l'estimation analytique lorsque "
            "la convergence avec le ratio et le retour d'expérience est jugée satisfaisante, sinon prévoir un arbitrage programmatique."
        )
    else:
        document.add_paragraph(
            "Recommandation MOE : utiliser l'estimation analytique comme base de chiffrage de référence, puis vérifier "
            "l'alignement des pièces écrites, bordereaux et ouvrages élémentaires avec les références comparables."
        )

    flux = BytesIO()
    document.save(flux)
    return flux.getvalue()


def _texte_clientele(clientele: str) -> str:
    return {
        "particulier_pme": "Particulier / petite PME",
        "public": "Maître d'ouvrage public",
        "cotraitrance": "Co-traitance",
        "sous_traitance": "Sous-traitance",
        "autre": "Autre contexte",
    }.get(clientele, clientele)


def _smic_horaire(localisation: str) -> Decimal:
    reference = _reference_sociale_localisation(localisation)
    if reference and getattr(reference, "smic_horaire", None):
        return Decimal(str(reference.smic_horaire))
    return SMIC_HORAIRE_2026.get(localisation or "metropole", SMIC_HORAIRE_2026["metropole"])


def _reference_sociale_localisation(localisation: str):
    if not localisation:
        return None
    from .models import ReferenceSocialeLocalisation

    return (
        ReferenceSocialeLocalisation.objects.filter(localisation=localisation, est_active=True)
        .order_by("-date_modification")
        .first()
    )


def _variante_locale_regle(regle, localisation: str):
    if not regle or not localisation:
        return None
    return getattr(regle, "variantes_locales", None).filter(localisation=localisation, est_active=True).first()


def _premiere_valeur(*valeurs):
    for valeur in valeurs:
        if valeur is not None and valeur != "":
            return valeur
    return None


def _questions_simulateur() -> list[dict[str, str]]:
    return [
        {"id": "remuneration", "libelle": "Quel est le niveau de rémunération retenu pour ce profil ?"},
        {"id": "contrat", "libelle": "S'agit-il d'un CDI ou d'un CDD, cadre ou non cadre, avec quels compléments employeur ?"},
        {"id": "production", "libelle": "Combien d'heures ou de jours sont réellement productifs sur l'année ?"},
        {"id": "structure", "libelle": "Quels coûts de structure, aléas, risque et marge faut-il intégrer ?"},
        {"id": "projection", "libelle": "Quel budget annuel et quel chiffre d'affaires cible faut-il viser pour ce recrutement ou ce profil ?"},
    ]


def _estimer_reduction_generale(
    remuneration_brute_reference: Decimal,
    charges_patronales_brutes: Decimal,
    smic_mensuel_reference: Decimal,
    appliquer_rgdu: bool,
) -> Decimal:
    if not appliquer_rgdu or smic_mensuel_reference <= 0 or remuneration_brute_reference <= 0:
        return Decimal("0")

    ratio = remuneration_brute_reference / smic_mensuel_reference
    if ratio >= Decimal("3"):
        return Decimal("0")

    coefficient_degressif = max(Decimal("0"), (Decimal("3") - ratio) / Decimal("2"))
    reduction = remuneration_brute_reference * Decimal("0.26") * coefficient_degressif
    return min(reduction, charges_patronales_brutes)


def donnees_simulation_depuis_profil(profil, surcharges: dict | None = None) -> dict:
    surcharges = surcharges or {}
    regle = getattr(profil, "regle_conventionnelle", None)
    convention = getattr(profil, "convention_collective", None)
    localisation = surcharges.get("localisation", profil.localisation)
    reference_locale = _reference_sociale_localisation(localisation)
    variante_locale = _variante_locale_regle(regle, localisation)
    libelle_localisation = LIBELLES_LOCALISATIONS.get(localisation, localisation)
    return {
        "profil_code": profil.code,
        "profil_libelle": profil.libelle,
        "localisation": localisation,
        "localisation_libelle": libelle_localisation,
        "convention_collective_code": surcharges.get("convention_collective_code", getattr(convention, "code", "")),
        "convention_collective_libelle": surcharges.get("convention_collective_libelle", getattr(convention, "libelle", "")),
        "regle_conventionnelle_code": surcharges.get("regle_conventionnelle_code", getattr(regle, "code", "")),
        "regle_conventionnelle_libelle": surcharges.get("regle_conventionnelle_libelle", getattr(regle, "libelle", "")),
        "variante_locale_regle_libelle": surcharges.get(
            "variante_locale_regle_libelle",
            getattr(variante_locale, "libelle", "") or (f"{getattr(regle, 'libelle', '')} — {libelle_localisation}" if variante_locale else ""),
        ),
        "reference_sociale_localisation_libelle": surcharges.get(
            "reference_sociale_localisation_libelle",
            getattr(reference_locale, "libelle", ""),
        ),
        "reference_sociale_source_officielle": surcharges.get(
            "reference_sociale_source_officielle",
            getattr(reference_locale, "source_officielle", ""),
        ),
        "commentaire_reglementaire_localisation": surcharges.get(
            "commentaire_reglementaire_localisation",
            getattr(reference_locale, "commentaire_reglementaire", ""),
        ),
        "commentaire_variante_locale": surcharges.get(
            "commentaire_variante_locale",
            getattr(variante_locale, "observations", ""),
        ),
        "salaire_brut_minimum_conventionnel": surcharges.get(
            "salaire_brut_minimum_conventionnel",
            _premiere_valeur(
                getattr(variante_locale, "salaire_brut_minimum_mensuel", None),
                getattr(regle, "salaire_brut_minimum_mensuel", None),
                Decimal("0"),
            ),
        ),
        "salaire_brut_mensuel": surcharges.get("salaire_brut_mensuel", profil.salaire_brut_mensuel_defaut),
        "primes_mensuelles": surcharges.get("primes_mensuelles", profil.primes_mensuelles_defaut),
        "avantages_mensuels": surcharges.get("avantages_mensuels", profil.avantages_mensuels_defaut),
        "heures_contractuelles_mensuelles": surcharges.get(
            "heures_contractuelles_mensuelles",
            _premiere_valeur(
                getattr(variante_locale, "heures_contractuelles_mensuelles_defaut", None),
                getattr(reference_locale, "heures_legales_mensuelles", None),
                getattr(regle, "heures_contractuelles_mensuelles_defaut", None),
                profil.heures_contractuelles_mensuelles,
            ),
        ),
        "heures_par_jour": surcharges.get(
            "heures_par_jour",
            _premiere_valeur(
                getattr(variante_locale, "heures_par_jour_defaut", None),
                getattr(regle, "heures_par_jour_defaut", None),
                profil.heures_par_jour,
            ),
        ),
        "taux_charges_salariales": surcharges.get(
            "taux_charges_salariales",
            _premiere_valeur(
                getattr(variante_locale, "taux_charges_salariales_defaut", None),
                profil.taux_charges_salariales,
            ),
        ),
        "taux_charges_patronales": surcharges.get(
            "taux_charges_patronales",
            _premiere_valeur(
                getattr(variante_locale, "taux_charges_patronales_defaut", None),
                profil.taux_charges_patronales,
            ),
        ),
        "taux_absenteisme": surcharges.get(
            "taux_absenteisme",
            _premiere_valeur(
                getattr(variante_locale, "taux_absenteisme_defaut", None),
                getattr(regle, "taux_absenteisme_defaut", None),
                profil.taux_absenteisme,
            ),
        ),
        "taux_temps_improductif": surcharges.get(
            "taux_temps_improductif",
            _premiere_valeur(
                getattr(variante_locale, "taux_temps_improductif_defaut", None),
                getattr(regle, "taux_temps_improductif_defaut", None),
                profil.taux_temps_improductif,
            ),
        ),
        "taux_frais_agence": surcharges.get("taux_frais_agence", profil.taux_frais_agence),
        "taux_risque_operationnel": surcharges.get("taux_risque_operationnel", profil.taux_risque_operationnel),
        "taux_marge_cible": surcharges.get("taux_marge_cible", profil.taux_marge_cible),
        "cout_equipement_mensuel": surcharges.get("cout_equipement_mensuel", profil.cout_equipement_mensuel),
        "cout_transport_mensuel": surcharges.get("cout_transport_mensuel", profil.cout_transport_mensuel),
        "cout_structure_mensuel": surcharges.get("cout_structure_mensuel", profil.cout_structure_mensuel),
        "clientele": surcharges.get("clientele", "public"),
        "contrat_travail": surcharges.get("contrat_travail", "cdi"),
        "statut_cadre": surcharges.get("statut_cadre", profil.categorie in {"ingenieur", "conducteur"}),
        "quotite_travail": surcharges.get("quotite_travail", Decimal("1")),
        "heures_supplementaires_mensuelles": surcharges.get("heures_supplementaires_mensuelles", Decimal("0")),
        "majoration_heures_supplementaires": surcharges.get("majoration_heures_supplementaires", Decimal("0.25")),
        "mutuelle_employeur_mensuelle": surcharges.get(
            "mutuelle_employeur_mensuelle",
            _premiere_valeur(
                getattr(variante_locale, "mutuelle_employeur_mensuelle_defaut", None),
                getattr(regle, "mutuelle_employeur_mensuelle_defaut", None),
                Decimal("55"),
            ),
        ),
        "titres_restaurant_employeur_mensuels": surcharges.get(
            "titres_restaurant_employeur_mensuels",
            _premiere_valeur(
                getattr(variante_locale, "titres_restaurant_employeur_mensuels_defaut", None),
                getattr(regle, "titres_restaurant_employeur_mensuels_defaut", None),
                Decimal("0"),
            ),
        ),
        "prime_transport_mensuelle": surcharges.get(
            "prime_transport_mensuelle",
            _premiere_valeur(
                getattr(variante_locale, "prime_transport_mensuelle_defaut", None),
                getattr(regle, "prime_transport_mensuelle_defaut", None),
                Decimal("0"),
            ),
        ),
        "appliquer_rgdu": surcharges.get("appliquer_rgdu", True),
        "taux_occupation_facturable": surcharges.get(
            "taux_occupation_facturable",
            _premiere_valeur(
                getattr(variante_locale, "taux_occupation_facturable_defaut", None),
                Decimal("0.78"),
            ),
        ),
        "jours_facturables_cibles_annuels": surcharges.get("jours_facturables_cibles_annuels"),
        "cout_recrutement_initial": surcharges.get(
            "cout_recrutement_initial",
            _premiere_valeur(
                getattr(variante_locale, "cout_recrutement_initial_defaut", None),
                getattr(regle, "cout_recrutement_initial_defaut", None),
                Decimal("0"),
            ),
        ),
        "nombre_profils": surcharges.get("nombre_profils", 1),
    }


def calculer_taux_horaire_reference_profil(profil, surcharges: dict | None = None) -> Decimal:
    """Retourne le taux horaire entreprise recommandé à partir d'un profil paramétré."""
    simulation = calculer_simulation_main_oeuvre(
        donnees_simulation_depuis_profil(profil, surcharges=surcharges)
    )
    return Decimal(str(simulation["resultats"]["taux_horaire_entreprise"]))


def calculer_simulation_main_oeuvre(donnees: dict) -> dict:
    clientele = donnees.get("clientele", "public")
    localisation = donnees.get("localisation", "metropole")
    contrat_travail = donnees.get("contrat_travail", "cdi")
    statut_cadre = bool(donnees.get("statut_cadre", False))
    quotite_travail = d(donnees.get("quotite_travail")) or Decimal("1")
    heures_supplementaires = d(donnees.get("heures_supplementaires_mensuelles"))
    majoration_heures_supp = d(donnees.get("majoration_heures_supplementaires")) or Decimal("0.25")
    mutuelle_employeur = d(donnees.get("mutuelle_employeur_mensuelle"))
    titres_restaurant = d(donnees.get("titres_restaurant_employeur_mensuels"))
    prime_transport = d(donnees.get("prime_transport_mensuelle"))
    appliquer_rgdu = bool(donnees.get("appliquer_rgdu", True))
    taux_occupation_facturable = d(donnees.get("taux_occupation_facturable")) or Decimal("0.78")
    jours_facturables_cibles_annuels = d(donnees.get("jours_facturables_cibles_annuels"))
    cout_recrutement_initial = d(donnees.get("cout_recrutement_initial"))
    nombre_profils = max(int(donnees.get("nombre_profils", 1) or 1), 1)
    salaire_brut_minimum_conventionnel = d(donnees.get("salaire_brut_minimum_conventionnel"))
    convention_collective_libelle = donnees.get("convention_collective_libelle", "")
    regle_conventionnelle_libelle = donnees.get("regle_conventionnelle_libelle", "")
    variante_locale_regle_libelle = donnees.get("variante_locale_regle_libelle", "")
    reference_sociale_localisation_libelle = donnees.get("reference_sociale_localisation_libelle", "")
    reference_sociale_source_officielle = donnees.get("reference_sociale_source_officielle", "")
    commentaire_reglementaire_localisation = donnees.get("commentaire_reglementaire_localisation", "")
    commentaire_variante_locale = donnees.get("commentaire_variante_locale", "")

    brut = d(donnees.get("salaire_brut_mensuel"))
    primes = d(donnees.get("primes_mensuelles"))
    avantages = d(donnees.get("avantages_mensuels"))
    heures_mensuelles = d(donnees.get("heures_contractuelles_mensuelles")) or Decimal("151.67")
    heures_mensuelles *= quotite_travail
    heures_par_jour = d(donnees.get("heures_par_jour")) or Decimal("7")

    taux_charges_salariales = d(donnees.get("taux_charges_salariales"))
    taux_charges_patronales = d(donnees.get("taux_charges_patronales"))
    taux_absenteisme = d(donnees.get("taux_absenteisme"))
    taux_temps_improductif = d(donnees.get("taux_temps_improductif"))
    taux_frais_agence = d(donnees.get("taux_frais_agence"))
    taux_risque = d(donnees.get("taux_risque_operationnel"))
    taux_marge = d(donnees.get("taux_marge_cible"))

    cout_equipement = d(donnees.get("cout_equipement_mensuel"))
    cout_transport = d(donnees.get("cout_transport_mensuel"))
    cout_structure = d(donnees.get("cout_structure_mensuel"))

    smic_horaire = _smic_horaire(localisation)
    smic_mensuel = arrondir(smic_horaire * heures_mensuelles)
    taux_horaire_brut_reference = brut / heures_mensuelles if heures_mensuelles > 0 else Decimal("0")
    montant_heures_supp = heures_supplementaires * taux_horaire_brut_reference * (Decimal("1") + majoration_heures_supp)
    remuneration_brute = brut + primes + avantages + montant_heures_supp
    cotisations_salariales = remuneration_brute * taux_charges_salariales
    net_hors_impot = remuneration_brute - cotisations_salariales
    charges_patronales_brutes = remuneration_brute * taux_charges_patronales
    reduction_generale = _estimer_reduction_generale(remuneration_brute, charges_patronales_brutes, smic_mensuel, appliquer_rgdu)
    charges_patronales = charges_patronales_brutes - reduction_generale

    cout_social_complements = mutuelle_employeur + titres_restaurant + prime_transport
    indemnite_precarite = remuneration_brute * Decimal("0.10") if contrat_travail == "cdd" else Decimal("0")
    cout_employeur = remuneration_brute + charges_patronales + cout_social_complements + indemnite_precarite
    couts_indirects = cout_equipement + cout_transport + cout_structure
    cout_complet_mensuel = cout_employeur + couts_indirects

    heures_payees_annuelles = heures_mensuelles * Decimal("12")
    heures_productives_annuelles = heures_payees_annuelles * (Decimal("1") - taux_absenteisme) * (
        Decimal("1") - taux_temps_improductif
    )
    if heures_productives_annuelles <= 0:
        heures_productives_annuelles = Decimal("1")

    cout_complet_annuel = cout_complet_mensuel * Decimal("12")
    cout_horaire_productif = cout_complet_annuel / heures_productives_annuelles
    cout_journalier_productif = cout_horaire_productif * heures_par_jour
    jours_productifs_annuels = heures_productives_annuelles / heures_par_jour if heures_par_jour > 0 else Decimal("0")
    jours_facturables = jours_facturables_cibles_annuels or (jours_productifs_annuels * taux_occupation_facturable)

    base_horaire_brut = remuneration_brute / heures_mensuelles if heures_mensuelles > 0 else Decimal("0")
    k_social = cout_employeur / remuneration_brute if remuneration_brute > 0 else Decimal("1")
    k_productivite = heures_payees_annuelles / heures_productives_annuelles if heures_productives_annuelles > 0 else Decimal("1")
    k_structure = Decimal("1") + taux_frais_agence + (couts_indirects / cout_employeur if cout_employeur > 0 else Decimal("0"))
    k_risque = Decimal("1") + taux_risque
    k_marge = Decimal("1") + taux_marge
    coefficient_k = k_social * k_productivite * k_structure * k_risque * k_marge

    taux_horaire_entreprise = base_horaire_brut * coefficient_k
    taux_journalier_entreprise = taux_horaire_entreprise * heures_par_jour
    chiffre_affaires_cible_annuel = taux_journalier_entreprise * jours_facturables
    marge_previsionnelle_annuelle = chiffre_affaires_cible_annuel - cout_complet_annuel - cout_recrutement_initial
    cout_total_premiere_annee = cout_complet_annuel + cout_recrutement_initial

    avertissements: list[str] = []
    minimum_reference = max(smic_mensuel, salaire_brut_minimum_conventionnel)
    if brut < smic_mensuel:
        avertissements.append(
            "Le salaire brut mensuel saisi est inférieur au Smic de référence de la localisation retenue."
        )
    if salaire_brut_minimum_conventionnel > 0 and brut < salaire_brut_minimum_conventionnel:
        avertissements.append(
            "Le salaire brut mensuel saisi est inférieur au minimum conventionnel paramétré pour ce profil."
        )
    if contrat_travail == "cdd":
        avertissements.append(
            "Le calcul intègre une indemnité de précarité estimative de 10 % du brut mensuel. Les cas d'exonération restent à vérifier."
        )
    if statut_cadre:
        avertissements.append(
            "Le statut cadre peut entraîner des écarts selon la convention collective et les garanties complémentaires de l'entreprise."
        )
    if localisation == "mayotte":
        avertissements.append(
            "La localisation Mayotte peut impliquer des paramètres sociaux et des dispositifs spécifiques : vérifiez les réglages territoriaux de la plateforme."
        )
    avertissements.append(
        "Simulation indicative : les conventions collectives, exonérations spécifiques et aides ciblées ne sont pas intégrées automatiquement."
    )

    return {
        "clientele": clientele,
        "clientele_libelle": _texte_clientele(clientele),
        "profil_libelle": donnees.get("profil_libelle") or donnees.get("profil_code") or "Profil",
        "convention_collective_libelle": convention_collective_libelle,
        "regle_conventionnelle_libelle": regle_conventionnelle_libelle,
        "variante_locale_regle_libelle": variante_locale_regle_libelle,
        "reference_sociale_localisation_libelle": reference_sociale_localisation_libelle,
        "localisation": localisation,
        "localisation_libelle": LIBELLES_LOCALISATIONS.get(localisation, localisation),
        "contrat_travail": contrat_travail,
        "statut_cadre": statut_cadre,
        "bulletin": {
            "salaire_brut_mensuel": float(arrondir(brut)),
            "primes_mensuelles": float(arrondir(primes)),
            "avantages_mensuels": float(arrondir(avantages)),
            "heures_supplementaires_mensuelles": float(arrondir(heures_supplementaires, 2)),
            "montant_heures_supplementaires": float(arrondir(montant_heures_supp)),
            "remuneration_brute_mensuelle": float(arrondir(remuneration_brute)),
            "cotisations_salariales": float(arrondir(cotisations_salariales)),
            "net_hors_impot": float(arrondir(net_hors_impot)),
            "charges_patronales_brutes": float(arrondir(charges_patronales_brutes)),
            "reduction_generale_estimee": float(arrondir(reduction_generale)),
            "charges_patronales": float(arrondir(charges_patronales)),
            "mutuelle_employeur_mensuelle": float(arrondir(mutuelle_employeur)),
            "titres_restaurant_employeur_mensuels": float(arrondir(titres_restaurant)),
            "prime_transport_mensuelle": float(arrondir(prime_transport)),
            "indemnite_precarite_mensuelle": float(arrondir(indemnite_precarite)),
            "cout_employeur_mensuel": float(arrondir(cout_employeur)),
            "couts_indirects_mensuels": float(arrondir(couts_indirects)),
            "cout_complet_mensuel": float(arrondir(cout_complet_mensuel)),
        },
        "production": {
            "heures_contractuelles_mensuelles": float(arrondir(heures_mensuelles)),
            "heures_payees_annuelles": float(arrondir(heures_payees_annuelles)),
            "heures_productives_annuelles": float(arrondir(heures_productives_annuelles)),
            "jours_productifs_annuels": float(arrondir(jours_productifs_annuels, 2)),
            "jours_facturables_cibles_annuels": float(arrondir(jours_facturables, 2)),
            "heures_par_jour": float(arrondir(heures_par_jour)),
            "quotite_travail": float(arrondir(quotite_travail, 4)),
            "taux_occupation_facturable": float(arrondir(taux_occupation_facturable, 4)),
            "taux_absenteisme": float(arrondir(taux_absenteisme, 4)),
            "taux_temps_improductif": float(arrondir(taux_temps_improductif, 4)),
        },
        "coefficients": {
            "k_social": float(arrondir(k_social, 4)),
            "k_productivite": float(arrondir(k_productivite, 4)),
            "k_structure": float(arrondir(k_structure, 4)),
            "k_risque": float(arrondir(k_risque, 4)),
            "k_marge": float(arrondir(k_marge, 4)),
            "coefficient_k_global": float(arrondir(coefficient_k, 4)),
        },
        "resultats": {
            "cout_horaire_productif": float(arrondir(cout_horaire_productif, 4)),
            "cout_journalier_productif": float(arrondir(cout_journalier_productif, 2)),
            "taux_horaire_entreprise": float(arrondir(taux_horaire_entreprise, 4)),
            "taux_journalier_entreprise": float(arrondir(taux_journalier_entreprise, 2)),
        },
        "projection_annuelle": {
            "cout_complet_annuel": float(arrondir(cout_complet_annuel)),
            "cout_recrutement_initial": float(arrondir(cout_recrutement_initial)),
            "cout_total_premiere_annee": float(arrondir(cout_total_premiere_annee)),
            "chiffre_affaires_cible_annuel": float(arrondir(chiffre_affaires_cible_annuel)),
            "marge_previsionnelle_annuelle": float(arrondir(marge_previsionnelle_annuelle)),
            "nombre_profils": nombre_profils,
            "cout_total_equipe_annuel": float(arrondir(cout_total_premiere_annee * Decimal(nombre_profils))),
            "chiffre_affaires_cible_equipe_annuel": float(arrondir(chiffre_affaires_cible_annuel * Decimal(nombre_profils))),
        },
        "hypotheses_reglementaires": {
            "source_smic": "Code du travail numérique — 01/2026",
            "smic_horaire_reference": float(arrondir(smic_horaire, 2)),
            "smic_mensuel_reference": float(arrondir(smic_mensuel)),
            "minimum_conventionnel_reference": float(arrondir(minimum_reference)),
            "reference_sociale_localisation_libelle": reference_sociale_localisation_libelle,
            "reference_sociale_source_officielle": reference_sociale_source_officielle,
            "commentaire_reglementaire_localisation": commentaire_reglementaire_localisation,
            "commentaire_variante_locale": commentaire_variante_locale,
            "appliquer_rgdu": appliquer_rgdu,
        },
        "avertissements": avertissements,
        "questions": _questions_simulateur(),
    }


def simuler_plan_activite(donnees: dict) -> dict:
    lignes_resultat = []
    total_cout = Decimal("0")
    total_ca = Decimal("0")
    total_marge = Decimal("0")

    for index, ligne in enumerate(donnees.get("lignes", []), start=1):
        simulation = calculer_simulation_main_oeuvre(ligne)
        effectif = max(int(ligne.get("effectif", 1) or 1), 1)
        projection = simulation["projection_annuelle"]
        cout_annuel_unitaire = d(projection["cout_total_premiere_annee"])
        ca_annuel_unitaire = d(projection["chiffre_affaires_cible_annuel"])
        marge_annuelle_unitaire = d(projection["marge_previsionnelle_annuelle"])

        cout_ligne = cout_annuel_unitaire * Decimal(effectif)
        ca_ligne = ca_annuel_unitaire * Decimal(effectif)
        marge_ligne = marge_annuelle_unitaire * Decimal(effectif)

        total_cout += cout_ligne
        total_ca += ca_ligne
        total_marge += marge_ligne

        lignes_resultat.append(
            {
                "ordre": index,
                "profil_code": ligne.get("profil_code"),
                "profil_libelle": simulation["profil_libelle"],
                "clientele": simulation["clientele"],
                "clientele_libelle": simulation["clientele_libelle"],
                "effectif": effectif,
                "simulation": simulation,
                "cout_total_annuel": float(arrondir(cout_ligne)),
                "chiffre_affaires_cible_annuel": float(arrondir(ca_ligne)),
                "marge_previsionnelle_annuelle": float(arrondir(marge_ligne)),
            }
        )

    taux_marge = (total_marge / total_ca) if total_ca > 0 else Decimal("0")

    return {
        "lignes": lignes_resultat,
        "totaux": {
            "cout_total_annuel": float(arrondir(total_cout)),
            "chiffre_affaires_cible_annuel": float(arrondir(total_ca)),
            "marge_previsionnelle_annuelle": float(arrondir(total_marge)),
            "taux_marge_previsionnelle": float(arrondir(taux_marge, 4)),
        },
        "avertissements": [
            "Les projections d'activité restent indicatives et doivent être consolidées avec votre convention collective, vos contrats collectifs et vos charges réelles.",
        ],
    }


def generer_pdf_simulation_main_oeuvre(simulation: dict) -> bytes:
    bulletin = simulation["bulletin"]
    production = simulation["production"]
    coefficients = simulation["coefficients"]
    resultats = simulation["resultats"]
    document = DocumentWord()

    titre = document.add_paragraph()
    run_titre = titre.add_run("Fiche de simulation main-d'œuvre")
    run_titre.bold = True
    run_titre.font.size = Pt(16)

    sous_titre = document.add_paragraph(
        f"Profil : {simulation['profil_libelle']} · Clientèle : {simulation['clientele_libelle']}"
    )
    sous_titre.runs[0].font.size = Pt(9)

    sections = [
        (
            "Bulletin synthétique",
            [
                ("Salaire brut mensuel", f"{bulletin['salaire_brut_mensuel']:.2f} €"),
                ("Primes", f"{bulletin['primes_mensuelles']:.2f} €"),
                ("Avantages", f"{bulletin['avantages_mensuels']:.2f} €"),
                ("Cotisations salariales", f"{bulletin['cotisations_salariales']:.2f} €"),
                ("Net hors impôt", f"{bulletin['net_hors_impot']:.2f} €"),
                ("Charges patronales", f"{bulletin['charges_patronales']:.2f} €"),
                ("Coût employeur mensuel", f"{bulletin['cout_employeur_mensuel']:.2f} €"),
                ("Coût complet mensuel", f"{bulletin['cout_complet_mensuel']:.2f} €"),
            ],
        ),
        (
            "Productivité",
            [
                ("Heures mensuelles", f"{production['heures_contractuelles_mensuelles']:.2f} h"),
                ("Heures payées annuelles", f"{production['heures_payees_annuelles']:.2f} h"),
                ("Heures productives annuelles", f"{production['heures_productives_annuelles']:.2f} h"),
                ("Heures par jour", f"{production['heures_par_jour']:.2f} h"),
                ("Absentéisme", f"{production['taux_absenteisme'] * 100:.2f} %"),
                ("Temps improductif", f"{production['taux_temps_improductif'] * 100:.2f} %"),
            ],
        ),
        (
            "Décomposition du coefficient K",
            [
                ("K social", f"{coefficients['k_social']:.4f}"),
                ("K productivité", f"{coefficients['k_productivite']:.4f}"),
                ("K structure", f"{coefficients['k_structure']:.4f}"),
                ("K risque", f"{coefficients['k_risque']:.4f}"),
                ("K marge", f"{coefficients['k_marge']:.4f}"),
                ("K global", f"{coefficients['coefficient_k_global']:.4f}"),
            ],
        ),
        (
            "Taux recommandés",
            [
                ("Coût horaire productif", f"{resultats['cout_horaire_productif']:.4f} €"),
                ("Coût journalier productif", f"{resultats['cout_journalier_productif']:.2f} €"),
                ("Taux horaire entreprise", f"{resultats['taux_horaire_entreprise']:.4f} €"),
                ("Taux journalier entreprise", f"{resultats['taux_journalier_entreprise']:.2f} €"),
            ],
        ),
    ]

    for titre_section, lignes in sections:
        document.add_heading(titre_section, level=2)
        tableau = document.add_table(rows=0, cols=2)
        tableau.style = "Table Grid"
        for libelle, valeur in lignes:
            cellules = tableau.add_row().cells
            cellules[0].text = libelle
            cellules[1].text = valeur

    flux_docx = BytesIO()
    document.save(flux_docx)
    contenu_docx = flux_docx.getvalue()

    with tempfile.TemporaryDirectory(prefix="lbh-simu-mo-") as dossier_temp:
        chemin_docx = Path(dossier_temp) / "simulation-main-oeuvre.docx"
        chemin_pdf = Path(dossier_temp) / "simulation-main-oeuvre.pdf"
        chemin_docx.write_bytes(contenu_docx)

        resultat = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                dossier_temp,
                str(chemin_docx),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if resultat.returncode != 0 or not chemin_pdf.exists():
            detail = (resultat.stderr or resultat.stdout or "").strip()
            raise RuntimeError(detail or "Échec de conversion PDF via LibreOffice.")

        return chemin_pdf.read_bytes()
